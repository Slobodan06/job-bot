from __future__ import annotations

import os
import re
import subprocess
import tempfile
from io import BytesIO

import fitz
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

_BULLET_CHARS = r"\-•*–—\u2022\u25cf\u25cb\u25aa\u25e6\u00b7\u2219"
_BULLET_PREFIX_RE = re.compile(rf"^[{_BULLET_CHARS}]\s*")


def _all_docx_paragraphs(doc: Document) -> list:
    paragraphs = list(doc.paragraphs)
    seen = {id(p._element) for p in paragraphs}
    for p_el in doc.element.xpath(".//*[local-name()='txbxContent']//w:p"):
        if id(p_el) not in seen:
            seen.add(id(p_el))
            paragraphs.append(Paragraph(p_el, doc))
    return paragraphs


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx") or name.endswith(".doc"):
        return _from_word(data, name)
    if name.endswith(".txt") or name.endswith(".md"):
        text = data.decode("utf-8", errors="replace")
        if text.startswith("\ufeff"):
            text = text.lstrip("\ufeff")
        return text
    raise ValueError("Unsupported file type. Use PDF, DOC, DOCX, or TXT.")


def _from_pdf(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        parts: list[str] = []
        for page in doc:
            parts.append(page.get_text("text") or "")
        return "\n".join(parts).strip()
    finally:
        doc.close()


def _looks_like_docx(data: bytes) -> bool:
    return data[:2] == b"PK"


def _from_word(data: bytes, filename: str) -> str:
    if _looks_like_docx(data):
        return _from_docx(data)
    if filename.endswith(".doc"):
        converted = _from_legacy_doc(data)
        if converted:
            return converted
    raise ValueError(
        "Could not read this Word file. Save it as .docx or PDF and upload again."
    )


def _paragraph_line(paragraph) -> str | None:
    text = (paragraph.text or "").strip()
    if not text:
        return None
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is not None and p_pr.find(qn("w:numPr")) is not None:
        return f"- {text}"
    style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
    if re.search(r"list|bullet", style_name, re.I):
        return f"- {text}"
    if _BULLET_PREFIX_RE.match(text):
        return f"- {_BULLET_PREFIX_RE.sub('', text).strip()}"
    return text


def _from_docx(data: bytes) -> str:
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for paragraph in _all_docx_paragraphs(doc):
        line = _paragraph_line(paragraph)
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
        if parts and parts[-1]:
            parts.append("")
    return "\n".join(parts).strip()


def _from_legacy_doc(data: bytes) -> str | None:
    """Extract text from binary .doc via antiword or LibreOffice when available."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as handle:
        handle.write(data)
        path = handle.name
    try:
        antiword = _run_text_extractor(["antiword", path])
        if antiword:
            return antiword
        return _convert_doc_with_soffice(path)
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def _run_text_extractor(cmd: list[str]) -> str | None:
    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=45,
            check=False,
        )
    except (FileNotFoundError, subprocess.SubprocessError):
        return None
    if result.returncode != 0:
        return None
    text = (result.stdout or "").strip()
    return text or None


def _convert_doc_with_soffice(doc_path: str) -> str | None:
    for binary in ("soffice", "libreoffice"):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                result = subprocess.run(
                    [
                        binary,
                        "--headless",
                        "--convert-to",
                        "txt:Text",
                        "--outdir",
                        tmp,
                        doc_path,
                    ],
                    capture_output=True,
                    timeout=60,
                    check=False,
                )
                if result.returncode != 0:
                    continue
                stem = os.path.splitext(os.path.basename(doc_path))[0]
                txt_path = os.path.join(tmp, f"{stem}.txt")
                if os.path.isfile(txt_path):
                    with open(txt_path, encoding="utf-8", errors="replace") as handle:
                        text = handle.read().strip()
                        if text:
                            return text
        except (FileNotFoundError, subprocess.SubprocessError):
            continue
    return None
