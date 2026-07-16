"""Local resume section detection — no external APIs."""
from __future__ import annotations

import re
import html
import zipfile
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from io import BytesIO

from docx import Document

from app.services.docx_resume import (
    DocxResumeDocument,
    _all_document_paragraphs,
    _experience_text_from_table_row,
    _group_experience_paragraph_indices,
    _field_hyperlink_urls_from_xml,
    _is_bullet_paragraph,
    _is_bulletish_text,
    _line_from_paragraph,
    _looks_like_skill_category_line,
    _paragraph_text,
    parse_resume_from_docx,
)
from app.services.pdf_resume import has_experience_date_range, sanitize_for_pdf
from app.services.extract_text import extract_text_from_bytes
from app.services.pdf_resume import (
    is_experience_role_header_line,
    merge_profile_links_into_contact,
    primary_role_header_from_block,
    split_experience_line_blocks,
)
from app.services.sectionize import is_pure_section_header, parse_resume_sections

_BULLET_CHARS = r"\-•*–—\u2022\u25cf\u25cb\u25aa\u25e6\u00b7\u2219"
_BULLET_LINE_RE = re.compile(rf"^[{_BULLET_CHARS}]\s*")
_EDUCATION_LINE_RE = re.compile(
    r"\b(bachelor|master|doctor|ph\.?\s*d|b\.?\s*s|b\.?\s*a|m\.?\s*s|m\.?\s*a|"
    r"associate|diploma|degree|university|college|school|institute|instituto|"
    r"technician|technologist|bootcamp|academy|licen[cs]iatura|grado)\b",
    re.I,
)
_SKILL_HEADING_RE = re.compile(
    r"^(shopify|frontend|front\s*end|backend|back\s*end|api|apis|database|data|tools|"
    r"languages?|frameworks?|platforms?|devops|cloud|testing|project\s+management|"
    r"methodolog\w*|ai|ai\s*&\s*llm|llm|engineering\s+practices?)\s*:?$",
    re.I,
)
_ACCOMPLISHMENT_LINE_RE = re.compile(
    r"^(?:Spearheaded|Engineered|Developed|Designed|Built|Led|Managed|Created|Implemented|"
    r"Automated|Optimized|Delivered|Architected|Established|Migrated|Integrated|Collaborated|"
    r"Enhanced|Reduced|Increased|Improved|Streamlined|Mentored|Set up|Created)\b",
    re.I,
)
_ROLE_TITLE_RE = re.compile(
    r"\b(engineer|developer|manager|architect|consultant|analyst|designer|specialist|lead|"
    r"principal|scientist|programmer|administrator|director|coordinator)\b",
    re.I,
)
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_PHONE_RE = re.compile(r"(?:\+?\d[\d\s().-]{7,}\d)")
_URL_RE = re.compile(r"https?://\S+", re.I)
_EXPERIENCE_DATE_RANGE_RE = re.compile(
    r"\b(?:(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+|"
    r"\d{1,2}/\s*)?(?:19|20)\d{2}\s*[\u2013\u2014-]\s*"
    r"(?:(?:(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|"
    r"jul(?:y)?|aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\s+|"
    r"\d{1,2}/\s*)?(?:19|20)\d{2}|present|current)\b",
    re.I,
)


@dataclass(frozen=True)
class ResolvedDocxSections:
    """Canonical resume sections mapped to the correct smart-template fields."""

    contact: str
    professional_summary: str
    professional_experience: str
    skills: str
    education: str
    other: str
    work_experience_roles: tuple[WorkExperienceRole, ...]
    experience_layout: str


@dataclass(frozen=True)
class WorkExperienceRole:
    header: str
    company: str
    title: str
    location: str
    period: str
    bullets: tuple[str, ...]


@dataclass(frozen=True)
class ResumeSectionAnalysis:
    contact: str
    professional_summary: str
    professional_experience: str
    skills: str
    education: str
    other: str
    work_experience_roles: tuple[WorkExperienceRole, ...]
    role_count: int
    experience_layout: str
    sections_detected: tuple[str, ...]
    source_format: str


def _strip_bullet_prefix(line: str) -> str:
    return _BULLET_LINE_RE.sub("", (line or "").strip()).strip()


def _line_is_experience_bullet(line: str) -> bool:
    stripped = (line or "").strip()
    if not stripped:
        return False
    if _BULLET_LINE_RE.match(stripped):
        return True
    if is_experience_role_header_line(stripped):
        return False
    if len(stripped) >= 40 and (
        _ACCOMPLISHMENT_LINE_RE.match(stripped)
        or (stripped.rstrip().endswith(".") and len(stripped.split()) >= 8)
    ):
        return True
    return False


def _parse_role_header(header_line: str) -> tuple[str, str, str, str]:
    """Best-effort split of 'Title | Company | Location | dates' or 'Company | Title | ...'."""
    stripped = (header_line or "").strip()
    if not stripped:
        return "", "", "", ""

    period = ""
    rest = stripped
    date_match = re.search(
        r"(\d{1,2}/\s*\d{4}\s*[–\-—]\s*(?:\d{1,2}/\s*\d{4}|present|current))",
        stripped,
        re.I,
    )
    if date_match:
        period = date_match.group(1).strip()
        rest = stripped[: date_match.start()].strip(" \t|,")

    if "\t" in rest:
        left, right = rest.split("\t", 1)
        rest = left.strip()
        if not period and has_experience_date_range(right):
            period = right.strip()

    parts = [part.strip() for part in rest.split("|") if part.strip()]
    title = ""
    company = ""
    location = ""

    if not parts:
        return company, title, location, period

    if len(parts) == 1:
        if _ROLE_TITLE_RE.search(parts[0]):
            title = parts[0]
        else:
            company = parts[0]
        return company, title, location, period

    first, second = parts[0], parts[1]
    first_is_title = bool(_ROLE_TITLE_RE.search(first))
    second_is_title = bool(_ROLE_TITLE_RE.search(second))

    if first_is_title and not second_is_title:
        title, company = first, second
        location = parts[2] if len(parts) > 2 else ""
    elif second_is_title and not first_is_title:
        company, title = first, second
        location = parts[2] if len(parts) > 2 else ""
    elif first_is_title:
        title, company = first, second
        location = parts[2] if len(parts) > 2 else ""
    else:
        company, title = first, second
        location = parts[2] if len(parts) > 2 else ""

    return company, title, location, period


def _looks_like_role_location(value: str) -> bool:
    clean = re.sub(r"\s+", " ", (value or "").strip(" |,-"))
    if not clean or len(clean) > 80 or _ROLE_TITLE_RE.search(clean):
        return False
    if re.fullmatch(r"remote(?:,\s*united states|,\s*usa)?", clean, re.I):
        return True
    parts = [part.strip() for part in clean.split(",") if part.strip()]
    return 2 <= len(parts) <= 4 and all(len(part.split()) <= 5 for part in parts)


def _role_parts(value: str) -> list[str]:
    clean = _EXPERIENCE_DATE_RANGE_RE.sub(" ", value or "")
    return [
        re.sub(r"\s+", " ", part).strip(" |,-")
        for part in re.split(r"[|\t\r\n]+", clean)
        if re.sub(r"\s+", " ", part).strip(" |,-")
    ]


def normalize_work_experience_role(role: WorkExperienceRole) -> WorkExperienceRole:
    """Remove repeated title/date/location metadata without changing resume facts."""
    period = ""
    for value in (role.period, role.header, role.company, role.title, role.location):
        match = _EXPERIENCE_DATE_RANGE_RE.search(value or "")
        if match:
            period = re.sub(r"\s+", " ", match.group(0)).strip()
            break

    parts: list[str] = []
    for value in (role.company, role.title, role.location, role.header):
        for part in _role_parts(value):
            if part.casefold() not in {item.casefold() for item in parts}:
                parts.append(part)

    title = next((part for part in _role_parts(role.title) if _ROLE_TITLE_RE.search(part)), "")
    if not title:
        title = next((part for part in parts if _ROLE_TITLE_RE.search(part)), "")

    def without_repeated_title(value: str) -> str:
        clean = value
        if title:
            clean = re.sub(rf"(?:\s*[,|]\s*)?{re.escape(title)}\s*$", "", clean, flags=re.I)
        return clean.strip(" |,-")

    normalized_parts = [without_repeated_title(part) for part in parts]
    normalized_parts = [part for part in normalized_parts if part]
    location = next(
        (
            without_repeated_title(part)
            for part in _role_parts(role.location)
            if _looks_like_role_location(without_repeated_title(part))
        ),
        "",
    )
    if not location:
        location = next((part for part in normalized_parts if _looks_like_role_location(part)), "")

    company = ""
    company_candidates: list[str] = []
    for source in (role.company, role.header):
        company_candidates.extend(_role_parts(source))
    for candidate in company_candidates:
        candidate = without_repeated_title(candidate)
        if not candidate:
            continue
        if title and candidate.casefold() == title.casefold():
            continue
        if location and candidate.casefold() == location.casefold():
            continue
        if _looks_like_role_location(candidate) or _ROLE_TITLE_RE.search(candidate):
            continue
        company = candidate
        break

    header_parts = [part for part in (company, title, location, period) if part]
    return WorkExperienceRole(
        header=" | ".join(header_parts) or role.header,
        company=company,
        title=title,
        location=location,
        period=period,
        bullets=role.bullets,
    )


def _role_from_block(header: str, bullets: list[str]) -> WorkExperienceRole:
    company, title, location, period = _parse_role_header(header)
    return WorkExperienceRole(
        header=header.strip(),
        company=company,
        title=title,
        location=location,
        period=period,
        bullets=tuple(bullets),
    )


def _roles_from_experience_text(experience: str) -> tuple[WorkExperienceRole, ...]:
    roles: list[WorkExperienceRole] = []
    for block in split_experience_line_blocks(experience or ""):
        if not block:
            continue
        header = primary_role_header_from_block(block) or block[0]
        bullets: list[str] = []
        for line in block:
            stripped = line.strip()
            if not stripped or stripped == header.strip():
                continue
            if _line_is_experience_bullet(stripped):
                bullets.append(_strip_bullet_prefix(stripped))
        roles.append(_role_from_block(header, bullets))
    return tuple(roles)


def _is_profile_link_line(text: str) -> bool:
    stripped = (text or "").strip()
    if _URL_RE.search(stripped):
        return True
    lower = stripped.lower()
    return "github.com" in lower or "linkedin.com" in lower or "gitlab.com" in lower


def _is_skill_category_line(text: str) -> bool:
    if _is_profile_link_line(text):
        return False
    return _looks_like_skill_category_line(text)


def _is_short_skill_heading(text: str, style_name: str = "") -> bool:
    stripped = _strip_bullet_prefix(text).strip()
    if not stripped or len(stripped) > 45:
        return False
    if has_experience_date_range(stripped) or _EDUCATION_LINE_RE.search(stripped):
        return False
    if _SKILL_HEADING_RE.match(stripped):
        return True
    return bool(re.search(r"heading\s*[23]?", style_name or "", re.I)) and not _ROLE_TITLE_RE.search(stripped)


def _is_education_entry_line(text: str) -> bool:
    stripped = _strip_bullet_prefix(text).strip(" |")
    if not stripped or len(stripped) > 180:
        return False
    if _EMAIL_RE.search(stripped) or _URL_RE.search(stripped):
        return False
    if _EDUCATION_LINE_RE.search(stripped):
        return True
    parts = [part.strip() for part in stripped.split("|") if part.strip()]
    return len(parts) >= 2 and has_experience_date_range(stripped) and any(
        _EDUCATION_LINE_RE.search(part) for part in parts
    )


def _looks_like_education_institution(text: str) -> bool:
    stripped = _strip_bullet_prefix(text).strip()
    if not stripped or len(stripped) > 100:
        return False
    return bool(re.search(r"\b(university|college|poly(?:technic)?|institute|school|academy)\b", stripped, re.I))


def _is_experience_header_not_education(text: str) -> bool:
    stripped = (text or "").strip()
    return is_experience_role_header_line(stripped) and not _is_education_entry_line(stripped)


def _normalize_skill_line(text: str) -> str:
    line = _strip_bullet_prefix(text)
    line = re.sub(r"^[\-•*–—\u2022\u25cf\u00b7\u2219]+\s*", "", line).strip()
    return line


def _is_prose_summary_line(text: str) -> bool:
    stripped = (text or "").strip()
    if len(stripped) < 85:
        return False
    if _EMAIL_RE.search(stripped) or stripped.lower().startswith("http"):
        return False
    if _is_skill_category_line(stripped):
        return False
    if stripped.endswith(".") or "experience" in stripped.lower() or "years" in stripped.lower():
        return True
    return len(stripped.split()) >= 12


def _format_contact_block(lines: list[str]) -> str:
    if not lines:
        return ""
    name = lines[0]
    headline = ""
    detail_parts: list[str] = []
    links: list[str] = []
    for line in lines[1:]:
        lower_line = line.lower()
        if name and name.lower() in lower_line and "portfolio" in lower_line and not _URL_RE.search(line):
            continue
        line = line.replace("\t", "|")
        if _EMAIL_RE.search(line) or _PHONE_RE.search(line):
            detail_parts.append(re.sub(r"\s+", " ", line.replace("•", "|").replace("·", "|")).strip())
        elif _URL_RE.search(line):
            links.append(_URL_RE.search(line).group(0).rstrip(".,;)"))  # type: ignore[union-attr]
        elif not headline and ("|" in line or _ROLE_TITLE_RE.search(line)):
            headline = line
        elif not headline and len(line) < 80:
            headline = line
        else:
            detail_parts.append(line)
    out = [sanitize_for_pdf(name)]
    if headline:
        out.append(sanitize_for_pdf(headline))
    if detail_parts:
        merged: list[str] = []
        for part in detail_parts:
            merged.extend(p.strip() for p in part.split("|") if p.strip())
        out.append(" | ".join(merged))
    for link in links:
        out.append(link)
    return "\n".join(out).strip()


def recover_contact_block_from_docx(docx_bytes: bytes, existing_contact: str = "") -> str:
    """Recover contact facts split across Word runs, tables, headers, or text boxes."""
    fragments: list[str] = []
    relationship_urls: list[str] = []
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
            names = [
                name for name in archive.namelist()
                if re.match(r"word/(?:document|header\d*|footer\d*)\.xml$", name, re.I)
            ]
            for name in names:
                xml_bytes = archive.read(name)
                xml = xml_bytes.decode("utf-8", errors="ignore")
                relationship_urls.extend(_field_hyperlink_urls_from_xml(xml_bytes))
                for paragraph in re.split(r"</w:p>", xml):
                    text = "".join(
                        html.unescape(value)
                        for value in re.findall(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", paragraph, re.I | re.S)
                    )
                    text = re.sub(r"\s+", " ", text).strip()
                    if text:
                        fragments.append(text)
            for name in archive.namelist():
                if not re.match(r"word/(?:_rels/|.+/_rels/).+\.rels$", name, re.I):
                    continue
                try:
                    root = ET.fromstring(archive.read(name))
                except ET.ParseError:
                    continue
                for relationship in root:
                    target = (relationship.attrib.get("Target") or "").strip()
                    target_mode = (relationship.attrib.get("TargetMode") or "").lower()
                    if target_mode == "external" and re.match(r"https?://", target, re.I):
                        relationship_urls.append(target.rstrip(".,;)"))
    except (OSError, zipfile.BadZipFile):
        return existing_contact

    existing_lines = [line.strip() for line in (existing_contact or "").splitlines() if line.strip()]
    corpus = "\n".join(fragments)
    emails = list(dict.fromkeys(_EMAIL_RE.findall(corpus)))
    phone_candidates = re.findall(r"(?<!\d)(?:\+?\d[\d\s().-]{7,}\d)(?!\d)", corpus)
    phones: list[str] = []
    for raw in phone_candidates:
        clean = re.sub(r"\s+", " ", raw).strip(" .,-")
        digits = re.sub(r"\D", "", clean)
        if 10 <= len(digits) <= 15 and not re.fullmatch(r"(?:19|20)\d{2}(?:19|20)\d{2}", digits):
            phones.append(clean)
    phones = list(dict.fromkeys(phones))
    urls = list(dict.fromkeys([
        *relationship_urls,
        *(
            match.rstrip(".,;)")
            for match in re.findall(r"(?:https?://|www\.)[^\s<>{}\[\]]+", corpus, re.I)
        ),
    ]))
    urls.sort(key=lambda value: 0 if re.search(r"linkedin\.com|github\.com", value, re.I) else 1)
    locations: list[str] = []

    candidate_name = existing_lines[0] if existing_lines else ""
    for fragment in fragments[:3]:
        header_candidate = _EMAIL_RE.sub(" ", fragment)
        if candidate_name:
            header_candidate = re.sub(re.escape(candidate_name), " ", header_candidate, flags=re.I)
        header_candidate = re.sub(r"\s+", " ", header_candidate).strip(" |,-")
        parts = [part.strip() for part in header_candidate.split(",")]
        if (
            len(parts) == 3
            and all(re.fullmatch(r"[A-Za-z][A-Za-z .'-]{1,40}", part) for part in parts)
            and not re.search(r"\b(?:RAG|AI|LLM|NLP)\b", header_candidate, re.I)
        ):
            locations.append(header_candidate)
            break

    location_re = re.compile(
        r"\b([A-Za-z][A-Za-z .'-]{1,40},\s*(?:[A-Z]{2}|United States|USA|Canada))\b"
    )
    for fragment in fragments[:35]:
        for match in location_re.finditer(fragment):
            candidate = match.group(1).strip(" |,-")
            candidate = re.sub(r"^(?:Remote|Location)\s*[:|-]\s*", "", candidate, flags=re.I)
            if candidate and candidate not in locations:
                locations.append(candidate)
    recovered: list[str] = []
    for value in [*emails[:1], *phones[:1], *locations[:1], *urls[:4]]:
        if value and value.lower() not in {line.lower() for line in existing_lines + recovered}:
            recovered.append(value)
    return "\n".join([*existing_lines, *recovered]).strip()


def _scan_table_template_body(doc: Document) -> tuple[str, str, str, str]:
    """Contact, summary, skills, and other from body paragraphs (table-label templates)."""
    contact_lines: list[str] = []
    summary = ""
    skill_lines: list[str] = []
    other_lines: list[str] = []
    phase = "contact"

    for para in doc.paragraphs:
        text = _paragraph_text(para).strip()
        if not text:
            continue
        if phase == "contact":
            if _is_profile_link_line(text):
                contact_lines.append(text)
                continue
            if _is_skill_category_line(text):
                phase = "skills"
                skill_lines.append(_normalize_skill_line(text))
                continue
            if _is_prose_summary_line(text):
                summary = sanitize_for_pdf(text)
                phase = "after_summary"
                continue
            contact_lines.append(text)
            continue
        if phase == "after_summary":
            if _is_skill_category_line(text):
                phase = "skills"
                skill_lines.append(_normalize_skill_line(text))
            continue
        elif phase == "skills":
            if _is_skill_category_line(text):
                skill_lines.append(_normalize_skill_line(text))
            elif _is_bulletish_text(text):
                phase = "other"
                other_lines.append(_normalize_skill_line(text))
            continue
        elif phase == "other" and (_is_bulletish_text(text) or len(text) < 120):
            other_lines.append(_normalize_skill_line(text))

    return (
        _format_contact_block(contact_lines),
        summary,
        "\n".join(skill_lines).strip(),
        "\n".join(other_lines).strip(),
    )


def _scan_paragraph_resume_sections(doc: Document) -> tuple[str, str, str, str, str] | None:
    """Infer sections from paragraph order when a resume omits literal section labels."""
    paragraphs = _all_document_paragraphs(doc)
    items: list[tuple[int, str, str]] = []
    for idx, para in enumerate(paragraphs):
        raw_line = _line_from_paragraph(para) or _paragraph_text(para)
        if not (raw_line or "").strip():
            raw_line = para.text or ""
        line = raw_line.strip()
        if not line:
            continue
        style_name = para.style.name if para.style and para.style.name else ""
        items.append((idx, style_name, line))
    if not items:
        return None

    first_role_pos = next(
        (i for i, (_, _, line) in enumerate(items) if _is_experience_header_not_education(line)),
        None,
    )
    if first_role_pos is None:
        return None

    skill_pos = next(
        (
            i
            for i, (_, style, line) in enumerate(items[first_role_pos + 1 :], start=first_role_pos + 1)
            if _is_skill_category_line(line) or _is_short_skill_heading(line, style)
        ),
        None,
    )
    exp_scan_end = skill_pos if skill_pos is not None else len(items)
    education_pos = next(
        (
            i
            for i, (_, _, line) in enumerate(items[first_role_pos + 1 : exp_scan_end], start=first_role_pos + 1)
            if _is_education_entry_line(line)
        ),
        None,
    )
    if education_pos is not None and education_pos > first_role_pos + 1:
        previous_line = items[education_pos - 1][2]
        if _looks_like_education_institution(previous_line):
            education_pos -= 1

    contact_lines: list[str] = []
    summary_lines: list[str] = []
    for _, _, line in items[:first_role_pos]:
        if _is_prose_summary_line(line):
            summary_lines.append(line)
        else:
            contact_lines.append(line)

    exp_end = education_pos if education_pos is not None else exp_scan_end
    experience_lines = [line for _, _, line in items[first_role_pos:exp_end]]
    education_lines: list[str] = []
    if education_pos is not None:
        education_end = skill_pos if skill_pos is not None else len(items)
        education_lines = [line for _, _, line in items[education_pos:education_end]]
    skill_lines: list[str] = []
    if skill_pos is not None:
        skill_lines = [
            _normalize_skill_line(line)
            for _, _, line in items[skill_pos:]
            if not is_pure_section_header(_normalize_skill_line(line))
        ]

    if not education_lines and not summary_lines and not skill_lines:
        return None

    return (
        _format_contact_block(contact_lines),
        "\n".join(summary_lines).strip(),
        "\n".join(experience_lines).strip(),
        "\n".join(skill_lines).strip(),
        "\n".join(education_lines).strip(),
    )


def _role_is_education(role: WorkExperienceRole) -> bool:
    return _is_education_entry_line(role.header) or _is_education_entry_line(
        " | ".join(part for part in (role.company, role.title, role.period) if part)
    )


def _role_from_experience_lines(lines: list[str]) -> WorkExperienceRole:
    header_lines: list[str] = []
    bullets: list[str] = []
    period = ""
    location = ""
    for line in lines:
        stripped = line.strip()
        if _BULLET_LINE_RE.match(stripped):
            bullets.append(_strip_bullet_prefix(stripped))
            continue
        if has_experience_date_range(stripped):
            parts = [p.strip() for p in re.split(r"[\n\r]+", stripped) if p.strip()]
            period = parts[0]
            if len(parts) > 1 and not location:
                location = parts[-1]
            continue
        if (
            not bullets
            and period
            and len(stripped) < 80
            and not _EMAIL_RE.search(stripped)
            and not _URL_RE.search(stripped)
        ):
            location = stripped
            continue
        if not bullets:
            header_lines.append(stripped)
    header = " | ".join(header_lines[:2]) if header_lines else (header_lines[0] if header_lines else "")
    company, title, loc_from_header, period_from_header = _parse_role_header(header)
    if not period and period_from_header:
        period = period_from_header
    if not location and loc_from_header:
        location = loc_from_header
    if not title and header_lines:
        title = header_lines[0]
    if not company and len(header_lines) > 1:
        company = header_lines[1]
    return WorkExperienceRole(
        header=header or (title or company),
        company=company,
        title=title,
        location=location,
        period=period,
        bullets=tuple(bullets),
    )


def roles_from_docx_table(docx_bytes: bytes) -> tuple[WorkExperienceRole, ...]:
    doc = Document(BytesIO(docx_bytes))
    from app.services.docx_resume import _detect_table_layout

    layout = _detect_table_layout(doc)
    if not layout:
        return ()
    _, _, exp_rows = layout
    roles: list[WorkExperienceRole] = []
    for ref in exp_rows:
        lines = [ln.strip() for ln in _experience_text_from_table_row(doc, ref) if ln.strip()]
        if lines:
            roles.append(_role_from_experience_lines(lines))
    return tuple(roles)


def resolve_docx_sections(
    docx_bytes: bytes,
    *,
    doc: DocxResumeDocument | None = None,
) -> ResolvedDocxSections:
    """Map Word content into canonical sections for smart PDF templates."""
    docx_doc = doc or parse_resume_from_docx(docx_bytes)
    word_doc = Document(BytesIO(docx_bytes))
    layout = "table" if docx_doc.experience_table_rows else "paragraph"

    roles = roles_from_docx_table(docx_bytes)
    if not roles:
        roles = _roles_from_docx_structure(docx_bytes, docx_doc)
    if not roles:
        roles = _roles_from_experience_text(docx_doc.parsed.professional_experience)

    if docx_doc.experience_table_rows:
        contact, summary, skills, other = _scan_table_template_body(word_doc)
        education = (docx_doc.parsed.education or "").strip()
        if not summary or len(summary.split()) <= 4:
            fallback = (docx_doc.parsed.professional_summary or "").strip()
            if fallback and len(fallback.split()) > 4:
                summary = fallback
        if not skills:
            skills = (docx_doc.parsed.skills or "").strip()
        experience = (docx_doc.parsed.professional_experience or "").strip()
    else:
        parsed = docx_doc.parsed
        contact = (parsed.contact or "").strip()
        summary = (parsed.professional_summary or "").strip()
        skills = (parsed.skills or "").strip()
        education = (parsed.education or "").strip()
        other = (parsed.other or "").strip()
        experience = (parsed.professional_experience or "").strip()

        structural = _scan_paragraph_resume_sections(word_doc)
        if structural is not None:
            s_contact, s_summary, s_experience, s_skills, s_education = structural
            if s_contact and (
                not contact
                or _is_prose_summary_line(contact)
                or any(_is_prose_summary_line(line) for line in contact.splitlines())
            ):
                contact = s_contact
            if s_summary and (not summary or len(summary.split()) <= 4):
                summary = s_summary
            if s_experience and (s_education or any(_role_is_education(role) for role in roles)):
                experience = s_experience
                roles = _roles_from_experience_text(experience)
            if s_skills and (not skills or len(s_skills) > len(skills)):
                skills = s_skills
            if s_education and (
                not education
                or (
                    any(_role_is_education(role) for role in roles)
                    and len([line for line in s_education.splitlines() if line.strip()])
                    >= len([line for line in education.splitlines() if line.strip()])
                )
            ):
                education = s_education

    dropped_education_roles = tuple(role for role in roles if _role_is_education(role))
    clean_roles = tuple(
        normalize_work_experience_role(role)
        for role in roles
        if not _role_is_education(role)
    )
    if len(clean_roles) != len(roles):
        if not education:
            edu_lines = [role.header for role in dropped_education_roles]
            education = "\n".join(edu_lines).strip()
    roles = clean_roles

    contact = recover_contact_block_from_docx(docx_bytes, contact)
    return ResolvedDocxSections(
        contact=contact,
        professional_summary=summary,
        professional_experience=experience,
        skills=skills,
        education=education,
        other=other,
        work_experience_roles=roles,
        experience_layout=layout,
    )


def _roles_from_docx_structure(docx_bytes: bytes, doc: DocxResumeDocument) -> tuple[WorkExperienceRole, ...]:
    """Detect bullets from Word table rows or list paragraphs."""
    if doc.experience_table_rows:
        table_roles = roles_from_docx_table(docx_bytes)
        if table_roles:
            return table_roles

    document = Document(BytesIO(docx_bytes))
    paragraphs = _all_document_paragraphs(document)
    indices = doc.section_body_indices.get("professional_experience") or []
    if not indices:
        return _roles_from_experience_text(doc.parsed.professional_experience)

    groups = _group_experience_paragraph_indices(document, indices)
    roles: list[WorkExperienceRole] = []
    for group in groups:
        header = ""
        bullets: list[str] = []
        for idx in group:
            if idx >= len(paragraphs):
                continue
            para = paragraphs[idx]
            text = _paragraph_text(para).strip()
            if not text:
                continue
            if _is_bullet_paragraph(para) or _is_bulletish_text(text):
                line = _line_from_paragraph(para) or text
                bullets.append(_strip_bullet_prefix(line))
                continue
            if not header and (is_experience_role_header_line(text) or not bullets):
                header = text
                continue
            if _line_is_experience_bullet(text):
                bullets.append(_strip_bullet_prefix(text))
        if header or bullets:
            roles.append(_role_from_block(header or "Professional Experience", bullets))
    return tuple(roles)


def _analysis_from_parsed(
    parsed,
    *,
    roles: tuple[WorkExperienceRole, ...],
    role_count: int,
    layout: str,
    source_format: str,
) -> ResumeSectionAnalysis:
    detected: list[str] = []
    for key, value in (
        ("contact", parsed.contact),
        ("professional_summary", parsed.professional_summary),
        ("professional_experience", parsed.professional_experience),
        ("skills", parsed.skills),
        ("education", parsed.education),
        ("other", parsed.other),
    ):
        if (value or "").strip():
            detected.append(key)

    normalized_roles = tuple(normalize_work_experience_role(role) for role in roles)
    return ResumeSectionAnalysis(
        contact=parsed.contact,
        professional_summary=parsed.professional_summary,
        professional_experience=parsed.professional_experience,
        skills=parsed.skills,
        education=parsed.education,
        other=parsed.other,
        work_experience_roles=normalized_roles,
        role_count=role_count or len(normalized_roles),
        experience_layout=layout,
        sections_detected=tuple(detected),
        source_format=source_format,
    )


def analyze_resume_sections(docx_bytes: bytes) -> ResumeSectionAnalysis:
    """Detect resume sections and work-experience roles from a .docx file."""
    doc = parse_resume_from_docx(docx_bytes)
    resolved = resolve_docx_sections(docx_bytes, doc=doc)
    role_count = len(resolved.work_experience_roles) or doc.detected_role_count or 0
    return _analysis_from_parsed(
        resolved,
        roles=resolved.work_experience_roles,
        role_count=role_count,
        layout=resolved.experience_layout,
        source_format="docx",
    )


def analyze_resume_sections_from_pdf(pdf_bytes: bytes) -> ResumeSectionAnalysis:
    """Best-effort section detection from PDF plain text (less accurate than .docx)."""
    text = extract_text_from_bytes("resume.pdf", pdf_bytes)
    if not text.strip():
        raise ValueError("Could not extract text from this PDF.")
    parsed = parse_resume_sections(text)
    parsed.contact = merge_profile_links_into_contact(
        parsed.contact,
        text,
        pdf_bytes=pdf_bytes,
    )
    roles = _roles_from_experience_text(parsed.professional_experience)
    return _analysis_from_parsed(
        parsed,
        roles=roles,
        role_count=len(roles),
        layout="pdf_text",
        source_format="pdf",
    )


def analyze_resume_file(data: bytes, *, filename: str) -> ResumeSectionAnalysis:
    name = (filename or "resume.docx").lower()
    if name.endswith(".pdf"):
        return analyze_resume_sections_from_pdf(data)
    if name.endswith(".docx") or name.endswith(".doc"):
        if name.endswith(".doc"):
            text = extract_text_from_bytes(filename, data)
            parsed = parse_resume_sections(text)
            roles = _roles_from_experience_text(parsed.professional_experience)
            return _analysis_from_parsed(
                parsed,
                roles=roles,
                role_count=len(roles),
                layout="doc_text",
                source_format="doc",
            )
        return analyze_resume_sections(data)
    raise ValueError("Unsupported file type. Upload .docx or .pdf.")
