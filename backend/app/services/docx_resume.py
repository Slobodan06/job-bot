"""Structured resume parsing and in-place tailoring for Word (.docx) uploads."""
from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from copy import deepcopy
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.services.pdf_resume import (
    default_ats_bullets_per_role,
    effective_bullets_per_role,
    has_experience_date_range,
    is_experience_role_header_line,
    line_is_factual_contact,
    merge_experience_headers_with_bullets,
    partition_experience_bullets_by_role,
    primary_role_header_from_block,
    split_experience_line_blocks,
    _looks_like_job_title_line,
    _link_label_for,
    _looks_like_role_header_line,
)
from app.services.sectionize import (
    ParsedResume,
    _partition_education_and_other,
    _separate_misplaced_jobs_from_education,
    implicit_section_after_contact,
    is_pure_section_header,
    match_section_header,
    parse_resume_sections,
)

_TAILOR_HEADER_SECTIONS = (
    "professional_summary",
    "professional_experience",
    "skills",
    "education",
    "other",
)
_FROZEN_DOCX_SECTIONS = frozenset({"contact", "education", "other"})
_EDITABLE_DOCX_SECTIONS = frozenset({"professional_summary", "professional_experience", "skills"})
_BULLET_CHARS = r"\-•*–—\u2022\u25cf\u25cb\u25aa\u25e6\u00b7\u2219"
_BULLET_CHAR_RE = re.compile(rf"^[{_BULLET_CHARS}]\s*")
_SKILL_CATEGORY_RE = re.compile(
    r"^[\-•*–—\u2022]?\s*(language|frontend|backend|ai/?ml|database|testing|devops|cloud|practices|tools|"
    r"frameworks?|platforms?|methodolog\w*)\s*:",
    re.I,
)
_TABLE_SUMMARY_RE = re.compile(r"^(professional\s+)?summary$|^profile$", re.I)
_TABLE_SKILLS_RE = re.compile(r"^skills$|^technical\s+skills$", re.I)
_TABLE_WORK_EXP_RE = re.compile(r"work\s+experience|^experience$", re.I)
_TABLE_EDUCATION_RE = re.compile(r"^education$", re.I)


@dataclass(frozen=True)
class ExperienceRowRef:
    """One job row inside a Word table (company/title/date cells stay frozen)."""

    table_idx: int
    row_idx: int
    content_cols: tuple[int, ...] = (0, 1)
    role_header_para_idx: int | None = None


@dataclass
class DocxResumeDocument:
    """Parsed resume content plus paragraph indices for in-place Word edits."""

    parsed: ParsedResume
    plain_text: str
    section_header_indices: dict[str, int] = field(default_factory=dict)
    section_body_indices: dict[str, list[int]] = field(default_factory=dict)
    contact_paragraph_indices: list[int] = field(default_factory=list)
    experience_table_rows: list[ExperienceRowRef] = field(default_factory=list)
    detected_role_count: int | None = None
    experience_bullet_slots: list[int] = field(default_factory=list)


_BULLET_PREFIX_RE = re.compile(rf"^[{_BULLET_CHARS}]\s+")
_ROLE_START_RE = re.compile(
    r"\b\d{1,2}/\s*\d{4}\s*[–\-—]\s*(?:\d{1,2}/\s*\d{4}|present|current)\b",
    re.I,
)


def _looks_like_docx(data: bytes) -> bool:
    return data[:2] == b"PK"


def _document_has_textboxes(document_xml: str) -> bool:
    return "txbxContent" in document_xml or "v:textbox" in document_xml or "wps:txbx" in document_xml


def _all_document_paragraphs(doc: Document) -> list[Paragraph]:
    """Body/table paragraphs plus text-box paragraphs (Canva/PDF-export templates)."""
    paragraphs = list(doc.paragraphs)
    seen = {id(p._element) for p in paragraphs}
    for p_el in doc.element.xpath(".//*[local-name()='txbxContent']//w:p"):
        if id(p_el) not in seen:
            seen.add(id(p_el))
            paragraphs.append(Paragraph(p_el, doc))
    return paragraphs


def _paragraph_at(doc: Document, index: int) -> Paragraph:
    return _all_document_paragraphs(doc)[index]


def _paragraph_text(paragraph: Paragraph) -> str:
    parts: list[str] = []
    for run in paragraph.runs:
        parts.append(run.text or "")
    if parts:
        return "".join(parts)
    return (paragraph.text or "").strip()


def _is_list_paragraph(paragraph: Paragraph) -> bool:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    return p_pr.find(qn("w:numPr")) is not None


def _is_bullet_paragraph(paragraph: Paragraph) -> bool:
    text = _paragraph_text(paragraph).strip()
    if not text:
        return _is_list_paragraph(paragraph)
    return _is_list_paragraph(paragraph) or bool(_BULLET_PREFIX_RE.match(text))


def _line_from_paragraph(paragraph: Paragraph) -> str | None:
    text = _paragraph_text(paragraph).strip()
    if not text:
        return None
    style_name = paragraph.style.name if paragraph.style and paragraph.style.name else ""
    if _is_list_paragraph(paragraph) or re.search(r"list|bullet", style_name, re.I):
        if not _BULLET_PREFIX_RE.match(text):
            return f"- {text}"
    return text


def _table_lines(doc: Document) -> list[str]:
    lines: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
        if lines and lines[-1]:
            lines.append("")
    return lines


def _strip_bullet_body(text: str) -> str:
    """Remove all leading bullet markers from tailored or source text."""
    body = (text or "").strip()
    while body:
        stripped = _BULLET_PREFIX_RE.sub("", body).strip()
        stripped = _BULLET_CHAR_RE.sub("", stripped).strip()
        if stripped == body:
            break
        body = stripped
    return body


def _existing_paragraph_bullet_prefix(paragraph: Paragraph) -> str | None:
    """Return the literal bullet prefix already stored in paragraph text, if any."""
    text = _paragraph_text(paragraph).strip()
    match = _BULLET_CHAR_RE.match(text)
    if not match:
        return None
    return match.group(0).rstrip()


def _format_bullet_for_paragraph(paragraph: Paragraph, line: str) -> str:
    """
    Format one experience bullet for a target paragraph.
    - Word list style only (no char in text): body only — Word renders the dot.
    - Existing character bullet (●, •, -): keep that single marker, update body.
    - No bullet at all: insert one.
    """
    body = _strip_bullet_body(line)
    if not body:
        return ""

    existing_prefix = _existing_paragraph_bullet_prefix(paragraph)
    if existing_prefix is not None:
        sep = "" if existing_prefix.endswith(" ") else " "
        return f"{existing_prefix}{sep}{body}"

    if _is_list_paragraph(paragraph):
        return body

    return f"• {body}"


def _display_line_for_paragraph(line: str, paragraph: Paragraph) -> str:
    if _is_bullet_paragraph(paragraph):
        # Skill/category lines (e.g. "• Frontend: React") keep the leading marker.
        if _looks_like_skill_category_line(line) or (
            _looks_like_skill_category_line(_paragraph_text(paragraph))
        ):
            return line.strip()
        return _BULLET_PREFIX_RE.sub("", line).strip()
    return line


def _replace_paragraph_text_inplace(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text while keeping existing runs/styles (first run keeps formatting)."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _replace_paragraph_text_plain(paragraph: Paragraph, text: str) -> None:
    """Replace paragraph text without keyword bolding; clear explicit run emphasis."""
    if paragraph.runs:
        paragraph.runs[0].text = text
        paragraph.runs[0].bold = False
        paragraph.runs[0].italic = False
        for run in paragraph.runs[1:]:
            run.text = ""
            run.bold = False
            run.italic = False
    else:
        run = paragraph.add_run(text)
        run.bold = False
        run.italic = False
    _strip_paragraph_character_emphasis(paragraph)


def _strip_paragraph_character_emphasis(paragraph: Paragraph) -> None:
    """Remove bold/italic from paragraph mark and all runs (Word stores these in XML)."""
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is not None:
        r_pr = p_pr.find(qn("w:rPr"))
        if r_pr is not None:
            for tag in ("w:i", "w:iCs", "w:b", "w:bCs"):
                el = r_pr.find(qn(tag))
                if el is not None:
                    r_pr.remove(el)
    for run in paragraph.runs:
        run.bold = False
        run.italic = False
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is not None:
            for tag in ("w:i", "w:iCs", "w:b", "w:bCs"):
                el = r_pr.find(qn(tag))
                if el is not None:
                    r_pr.remove(el)


_HIGHLIGHT_METRIC_RE = re.compile(
    r"\d+(?:\.\d+)?\s*(?:%|percent\b)|"
    r"\b\d+\+\b|"
    r"\b\d{1,3}(?:,\d{3})+\+?\b|"
    r"\b(?:reduced|increased|improved|cut|decreased|accelerated|boosted|lowered|grew|saved|"
    r"delivered|processed|scaled|optimized)\b[^.!?\n]{0,40}\bby\s+\d+(?:\.\d+)?\s*(?:%|x\b|X\b)?",
    re.I,
)
_HIGHLIGHT_TECH_RE = re.compile(
    r"\b(?:React(?:\.js)?|Angular|Vue(?:\.js)?|Node(?:\.js)?|TypeScript|JavaScript|Python|"
    r"Java|Kotlin|Go\b|Golang|C#|\.NET|AWS|Azure|GCP|Docker|Kubernetes|Terraform|CI/CD|"
    r"PostgreSQL|MongoDB|Redis|GraphQL|REST(?:ful)?|APIs?|FastAPI|Django|Flask|Rails|"
    r"Next(?:\.js)?|Express(?:\.js)?|Kubernetes|Lambda|OpenAI|TensorFlow|PyTorch)\b",
    re.I,
)


def _copy_run_font(source_run, target_run) -> None:
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def _bold_spans_for_text(
    text: str,
    highlight_terms: list[str] | None,
    *,
    auto_tech_and_metrics: bool = True,
    auto_metrics: bool = False,
) -> list[tuple[int, int]]:
    if not text:
        return []
    spans: list[tuple[int, int]] = []

    def _overlaps(start: int, end: int) -> bool:
        return any(not (end <= s or start >= e) for s, e in spans)

    if auto_tech_and_metrics or auto_metrics:
        for match in _HIGHLIGHT_METRIC_RE.finditer(text):
            if not _overlaps(match.start(), match.end()):
                spans.append((match.start(), match.end()))

    if auto_tech_and_metrics:
        for match in _HIGHLIGHT_TECH_RE.finditer(text):
            if len(match.group(0)) >= 3 and not _overlaps(match.start(), match.end()):
                spans.append((match.start(), match.end()))

    terms = sorted({t.strip() for t in (highlight_terms or []) if t and len(t.strip()) >= 3}, key=len, reverse=True)
    for term in terms:
        if len(term) > 80:
            continue
        if re.match(r"^[\w\-./+#]+$", term):
            pattern = re.compile(r"(?<![\w\-./+#])" + re.escape(term) + r"(?![\w\-./+#])", re.I)
        else:
            pattern = re.compile(re.escape(term), re.I)
        for match in pattern.finditer(text):
            if not _overlaps(match.start(), match.end()):
                spans.append((match.start(), match.end()))

    spans.sort()
    merged: list[tuple[int, int]] = []
    for start, end in spans:
        if merged and start <= merged[-1][1]:
            merged[-1] = (merged[-1][0], max(merged[-1][1], end))
        else:
            merged.append((start, end))
    return merged


def _replace_paragraph_text_with_highlights(
    paragraph: Paragraph,
    text: str,
    highlight_terms: list[str] | None = None,
    *,
    auto_tech_and_metrics: bool = True,
    auto_metrics: bool = False,
) -> None:
    """Replace paragraph text, bolding highlight terms (and optionally tech/metrics)."""
    spans = _bold_spans_for_text(
        text,
        highlight_terms,
        auto_tech_and_metrics=auto_tech_and_metrics,
        auto_metrics=auto_metrics,
    )
    if not spans:
        _replace_paragraph_text_inplace(paragraph, text)
        return

    template_run = paragraph.runs[0] if paragraph.runs else None
    p_element = paragraph._element
    for child in list(p_element):
        if child.tag == qn("w:r"):
            p_element.remove(child)

    pos = 0
    for start, end in spans:
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            if template_run:
                _copy_run_font(template_run, run)
        bold_run = paragraph.add_run(text[start:end])
        bold_run.bold = True
        if template_run:
            _copy_run_font(template_run, bold_run)
        pos = end
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if template_run:
            _copy_run_font(template_run, run)


def _sync_paragraph_runs(dup_para: Paragraph, src_para: Paragraph) -> None:
    if dup_para._element is src_para._element:
        return
    dup_el = dup_para._element
    for child in list(dup_el):
        if child.tag == qn("w:r"):
            dup_el.remove(child)
    for child in src_para._element:
        if child.tag == qn("w:r"):
            dup_el.append(deepcopy(child))


def _clone_and_insert_after(
    anchor: Paragraph,
    template: Paragraph,
    text: str,
    highlight_terms: list[str] | None = None,
) -> Paragraph:
    if _is_bullet_paragraph(template):
        display = _format_bullet_for_paragraph(template, text)
    else:
        display = _display_line_for_paragraph(text, template)
    new_p = deepcopy(template._element)
    anchor._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor._parent)
    _replace_paragraph_text_with_highlights(new_para, display, highlight_terms)
    return new_para


def _delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _delete_paragraphs_at_indices(doc: Document, indices: list[int]) -> None:
    for idx in sorted(set(indices), reverse=True):
        if 0 <= idx < len(doc.paragraphs):
            _delete_paragraph(doc.paragraphs[idx])


def _strip_paragraph_list_formatting(paragraph: Paragraph) -> None:
    """Remove Word list/numbering so an emptied paragraph does not render a bullet glyph."""
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return
    for tag in ("w:numPr", "w:ilvl"):
        el = p_pr.find(qn(tag))
        if el is not None:
            p_pr.remove(el)


def _collapse_paragraph_spacing(paragraph: Paragraph) -> None:
    """Minimize vertical space for emptied paragraphs (zero margins, 1twip line height)."""
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    try:
        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 0.01
    except Exception:
        pass
    p_pr = _paragraph_ppr(paragraph)
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "1")
    spacing.set(qn("w:lineRule"), "exact")
    _apply_vanish_to_paragraph(paragraph)


def _apply_vanish_to_paragraph(paragraph: Paragraph) -> None:
    """Mark paragraph runs hidden and tiny so empty slots take no visible space."""
    from docx.oxml import OxmlElement

    def _ensure_vanish_run(run) -> None:
        r_pr = run._element.find(qn("w:rPr"))
        if r_pr is None:
            r_pr = OxmlElement("w:rPr")
            run._element.insert(0, r_pr)
        for tag in ("w:vanish",):
            if r_pr.find(qn(tag)) is None:
                r_pr.append(OxmlElement(tag))
        sz = r_pr.find(qn("w:sz"))
        if sz is None:
            sz = OxmlElement("w:sz")
            r_pr.append(sz)
        sz.set(qn("w:val"), "2")

    if paragraph.runs:
        for run in paragraph.runs:
            _ensure_vanish_run(run)
    else:
        run = paragraph.add_run("")
        _ensure_vanish_run(run)


def _clear_paragraph_visible(paragraph: Paragraph) -> None:
    """Clear unused bullet-slot text and hide the slot (does not touch layout spacer paragraphs)."""
    _strip_paragraph_list_formatting(paragraph)
    _replace_paragraph_text_plain(paragraph, "")
    _collapse_paragraph_spacing(paragraph)
    # Ensure no ghost text remains for the XML merge pass (PDF converters ignore w:vanish).
    for run in paragraph.runs:
        run.text = ""


def _paragraph_has_layout_spacer_xml(paragraph_xml: str) -> bool:
    """True when an empty paragraph carries intentional vertical space (section/job gaps)."""
    if _paragraph_plain_text_from_xml(paragraph_xml).strip():
        return False
    if re.search(r'w:before="[1-9]\d*"', paragraph_xml):
        return True
    if re.search(r'w:after="[1-9]\d*"', paragraph_xml):
        return True
    if re.search(r'w:line="(?!240\b|276\b|259\b)[1-9]\d{2,}"', paragraph_xml):
        return True
    if "<w:br" in paragraph_xml:
        return True
    return False


def _strip_skill_category_label(text: str) -> str:
    return re.sub(rf"^[{_BULLET_CHARS}]\s*", "", (text or "").strip()).strip().rstrip(":")


def _parse_skills_category_map(skills_text: str) -> dict[str, list[str]]:
    """Map category label -> tool tokens from paired or colon-formatted skills text."""
    lines = [line.strip() for line in (skills_text or "").splitlines() if line.strip()]
    if not lines:
        return {}
    out: dict[str, list[str]] = {}
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        colon_match = re.match(r"^([^:]+):\s*(.+)$", _strip_skill_category_label(line))
        if colon_match and colon_match.group(2).strip():
            label = colon_match.group(1).strip()
            tokens = [t.strip() for t in colon_match.group(2).split(",") if t.strip()]
            out[label.lower()] = tokens
            idx += 1
            continue
        label = _strip_skill_category_label(line)
        tokens: list[str] = []
        if idx + 1 < len(lines):
            nxt = lines[idx + 1]
            if "," in nxt and not _strip_skill_category_label(nxt).endswith(":"):
                tokens = [t.strip() for t in nxt.split(",") if t.strip()]
                idx += 2
                out[label.lower()] = tokens
                continue
        out[label.lower()] = tokens
        idx += 1
    return out


def _skill_line_has_inline_tools(text: str) -> bool:
    """True for 'Category: tool, tool' on one bullet line (Ciro-style skills)."""
    stripped = _strip_skill_category_label(text or "")
    return bool(re.match(r"^[^:]+:\s*\S", stripped))


def _pair_skill_paragraph_indices(doc: Document, indices: list[int]) -> list[tuple[int, int | None]]:
    """Alternating (category_header_idx, tools_content_idx) slots in Cameron-style templates."""
    paragraphs = _all_document_paragraphs(doc)
    pairs: list[tuple[int, int | None]] = []
    idx = 0
    while idx < len(indices):
        header_idx = indices[idx]
        header_para = paragraphs[header_idx]
        header_text = _paragraph_text(header_para).strip()
        if _skill_line_has_inline_tools(header_text):
            pairs.append((header_idx, None))
            idx += 1
            continue
        content_idx = indices[idx + 1] if idx + 1 < len(indices) else None
        if content_idx is not None:
            content_para = paragraphs[content_idx]
            header_text = _paragraph_text(header_para).strip()
            content_text = _paragraph_text(content_para).strip()
            header_is_category = _is_bullet_paragraph(header_para) or (
                "," not in header_text and len(header_text) < 60
            )
            content_is_tools = bool(content_text) and (
                "," in content_text or not _is_bullet_paragraph(content_para)
            )
            if header_is_category and content_is_tools:
                pairs.append((header_idx, content_idx))
                idx += 2
                continue
        pairs.append((header_idx, content_idx))
        idx += 2 if content_idx is not None else 1
    return pairs


def _skill_content_indices(doc: Document, indices: list[int]) -> list[int]:
    editable: list[int] = []
    for header_idx, content_idx in _pair_skill_paragraph_indices(doc, indices):
        if content_idx is not None:
            editable.append(content_idx)
        else:
            editable.append(header_idx)
    return editable


def _update_skills_preserving_template(
    doc: Document,
    indices: list[int],
    new_text: str,
    source_text: str,
    *,
    highlight_terms: list[str] | None = None,
    enable_bold: bool = True,
) -> None:
    """
    Update only the tool-list paragraphs in a paired skills template.
    Category header paragraphs (bullets) are never modified — preserves Word UI.
    """
    if not indices:
        return
    paragraphs = _all_document_paragraphs(doc)
    tailored_map = _parse_skills_category_map(new_text)
    source_map = _parse_skills_category_map(source_text)
    pairs = _pair_skill_paragraph_indices(doc, indices)
    if not pairs:
        _update_lines_index_preserving(
            doc,
            indices,
            new_text,
            highlight_terms=highlight_terms,
            selective_highlight=True,
            enable_bold=enable_bold,
            clear_unused=False,
        )
        return

    for header_idx, content_idx in pairs:
        header_para = paragraphs[header_idx]
        label = _strip_skill_category_label(_paragraph_text(header_para))
        key = label.lower().rstrip(":")
        if key in tailored_map:
            tokens = tailored_map[key]
        else:
            tokens = source_map.get(key) or []
        if content_idx is None or content_idx >= len(paragraphs):
            category = label.rstrip(":").strip()
            line = f"{category}: {', '.join(tokens)}" if tokens else category
            display = (
                _format_bullet_for_paragraph(header_para, line)
                if _is_bullet_paragraph(header_para)
                else line
            )
            _apply_paragraph_text(
                header_para,
                display,
                highlight_terms=highlight_terms,
                plain=False,
                selective_highlight=True,
                enable_bold=enable_bold,
            )
            continue
        line = ", ".join(tokens)
        _apply_paragraph_text(
            paragraphs[content_idx],
            line,
            highlight_terms=highlight_terms,
            plain=False,
            selective_highlight=True,
            enable_bold=enable_bold,
        )


def _layout_protected_paragraph_indices(
    contact_paragraph_indices: list[int],
    section_body_indices: dict[str, list[int]],
    section_header_indices: dict[str, int],
    *,
    skill_header_indices: list[int] | None = None,
    doc: Document | None = None,
) -> set[int]:
    """Paragraphs whose spacing/XML must never be collapsed (header, spacers, skill labels)."""
    protected = set(contact_paragraph_indices)
    for idx in section_header_indices.values():
        if idx >= 0:
            protected.add(idx)
    skill_indices = section_body_indices.get("skills", [])
    if skill_indices:
        first_skill = min(skill_indices)
        for i in range(0, first_skill):
            protected.add(i)
    if skill_header_indices:
        protected.update(skill_header_indices)
    return protected


def _collapse_empty_paragraph_spacing(
    doc: Document,
    *,
    skip_indices: set[int] | None = None,
) -> None:
    """Collapse spacing only on empty paragraphs cleared during editing — never layout spacers."""
    skip = skip_indices or set()
    for idx, paragraph in enumerate(_all_document_paragraphs(doc)):
        if idx in skip:
            continue
        if not _paragraph_text(paragraph).strip():
            if _paragraph_has_layout_spacer_xml(paragraph._element.xml):
                continue
            _collapse_paragraph_spacing(paragraph)


def _apply_bullet_updates(
    slot_paragraphs: list[Paragraph],
    new_bullets: list[str],
    template_para: Paragraph,
    *,
    highlight_terms: list[str] | None = None,
    enable_bold: bool = True,
) -> None:
    """
    Update bullet slots in place only.
    Paragraph count must stay identical to the upload so frozen-docx XML merge stays aligned.
    Extra tailored bullets are truncated; unused slots are cleared without deleting paragraphs.
    """
    if not slot_paragraphs:
        return
    for bi, para in enumerate(slot_paragraphs):
        if bi < len(new_bullets):
            formatted = _format_bullet_for_paragraph(para, new_bullets[bi])
            _apply_paragraph_text(
                para,
                formatted,
                highlight_terms=highlight_terms,
                selective_highlight=True,
                bold_metrics=True,
                enable_bold=enable_bold,
            )
        else:
            _clear_paragraph_visible(para)


def _split_tailored_lines(text: str) -> list[str]:
    lines: list[str] = []
    for raw in (text or "").replace("\r\n", "\n").split("\n"):
        stripped = raw.strip()
        if stripped:
            lines.append(stripped)
    return lines


def _partition_header_and_bullet_lines(lines: list[str]) -> tuple[list[str], list[str]]:
    headers: list[str] = []
    bullets: list[str] = []
    for line in lines:
        if _BULLET_PREFIX_RE.match(line):
            bullets.append(line)
        else:
            headers.append(line)
    return headers, bullets


def _detect_contact_header_role_line(text: str) -> bool:
    """True for the headline job-title line under the candidate name (not email/phone/links)."""
    stripped = (text or "").strip()
    if not stripped or line_is_factual_contact(stripped):
        return False
    # Full experience headers (Company | Title | Location + dates) are not contact headlines.
    if has_experience_date_range(stripped):
        return False
    if stripped.count("|") >= 2:
        return False
    if "\t" in stripped and re.search(r"\d{1,2}/\s*\d{4}", stripped):
        return False
    if _ROLE_START_RE.search(stripped):
        return False
    if _looks_like_role_header_line(stripped) and "|" in stripped:
        return True
    if "|" in stripped and len(stripped) < 160:
        return bool(
            re.search(
                r"\b(engineer|developer|architect|consultant|manager|analyst|designer|specialist|lead|"
                r"principal|scientist|programmer|full[\s-]?stack|software|devops|sre)\b",
                stripped,
                re.I,
            )
        )
    if re.search(
        r"\b(engineer|developer|architect|consultant|manager|analyst|designer|specialist|lead|"
        r"principal|scientist|programmer|full[\s-]?stack|software|devops|sre)\b",
        stripped,
        re.I,
    ):
        return len(stripped) < 120
    return False


def _is_experience_role_start(paragraph: Paragraph, text: str) -> bool:
    if _is_bullet_paragraph(paragraph):
        return False
    if _detect_contact_header_role_line(text):
        return False
    stripped = (text or "").strip()
    if _looks_like_role_header_line(stripped):
        return True
    if line_is_factual_contact(text):
        return False
    # In-role sub-headings (e.g. "DevOps & Cloud Engineering:") are not new employers.
    if stripped.rstrip().endswith(":") and "|" not in stripped and not _ROLE_START_RE.search(stripped):
        return False
    if _ROLE_START_RE.search(stripped):
        return True
    if "|" in stripped and len(stripped) < 120:
        return True
    if re.search(
        r"\b(engineer|developer|manager|architect|analyst|consultant|specialist|director|lead|principal)\b",
        stripped,
        re.I,
    ):
        return len(stripped) < 100 and ("|" in stripped or "\t" in stripped or _ROLE_START_RE.search(stripped))
    return False


def _normalize_experience_header_key(line: str) -> str:
    return re.sub(r"\s+", " ", (line or "").strip().lower())


def _group_indices_by_text_blocks(
    paragraphs: list[Paragraph],
    indices: list[int],
    text_blocks: list[list[str]],
) -> list[list[int]]:
    """Map plain-text experience blocks back to Word paragraph indices."""
    if not indices or not text_blocks:
        return [indices] if indices else []

    header_keys = [
        _normalize_experience_header_key(primary_role_header_from_block(block) or (block[0] if block else ""))
        for block in text_blocks
    ]
    groups: list[list[int]] = [[] for _ in text_blocks]
    block_i = 0

    for idx in indices:
        text = _paragraph_text(paragraphs[idx]).strip()
        if not text:
            if groups[block_i]:
                groups[block_i].append(idx)
            continue

        norm = _normalize_experience_header_key(text)
        if block_i < len(header_keys) - 1 and norm == header_keys[block_i + 1]:
            block_i += 1
        elif (
            block_i < len(text_blocks) - 1
            and is_experience_role_header_line(text)
            and _looks_like_job_title_line(text)
            and norm == header_keys[block_i + 1]
        ):
            block_i += 1

        groups[block_i].append(idx)

    non_empty = [group for group in groups if group]
    return non_empty if non_empty else [indices]


def _block_has_bullets(paragraphs: list[Paragraph], indices: list[int]) -> bool:
    return any(_is_bullet_paragraph(paragraphs[i]) for i in indices if i < len(paragraphs))


def _group_experience_paragraph_indices(doc: Document, indices: list[int]) -> list[list[int]]:
    if not indices:
        return []
    paragraphs = _all_document_paragraphs(doc)
    blocks: list[list[int]] = []
    current: list[int] = []
    for idx in indices:
        para = paragraphs[idx]
        text = _paragraph_text(para).strip()
        if current and text:
            starts_new_role = _is_experience_role_start(para, text)
            if not starts_new_role and _block_has_bullets(paragraphs, current):
                if is_experience_role_header_line(text) or (
                    _looks_like_job_title_line(text) and not _is_bullet_paragraph(para)
                ):
                    starts_new_role = True
            if starts_new_role:
                blocks.append(current)
                current = [idx]
                continue
        current.append(idx)
    if current:
        blocks.append(current)

    experience_lines: list[str] = []
    for idx in indices:
        line = _line_from_paragraph(paragraphs[idx])
        if line:
            experience_lines.append(line)
    text_blocks = split_experience_line_blocks("\n".join(experience_lines))
    if len(text_blocks) > len(blocks):
        return _group_indices_by_text_blocks(paragraphs, indices, text_blocks)
    return blocks if blocks else [indices]


def _experience_template_bullet_slots(doc: Document, indices: list[int]) -> list[int]:
    """Bullet paragraph slots per role block in the uploaded template (caps export capacity)."""
    paragraphs = _all_document_paragraphs(doc)
    blocks = _group_experience_paragraph_indices(doc, indices)
    return [
        sum(1 for idx in block if _is_bullet_paragraph(paragraphs[idx]))
        for block in blocks
    ]


def _resolve_experience_bullet_targets(
    doc: Document,
    indices: list[int],
    bullets_per_role: list[int] | None,
) -> list[int]:
    """ATS bullet targets capped to template slots with a 3-bullet floor per employer."""
    blocks = _group_experience_paragraph_indices(doc, indices)
    role_count = len(blocks)
    if role_count <= 0:
        return []
    ats_targets = (
        bullets_per_role
        if bullets_per_role and len(bullets_per_role) == role_count
        else default_ats_bullets_per_role(role_count)
    )
    template_slots = _experience_template_bullet_slots(doc, indices)
    return effective_bullets_per_role(ats_targets, template_slots)


def _bullet_lines(text: str) -> list[str]:
    return [line for line in _split_tailored_lines(text) if _BULLET_PREFIX_RE.match(line)]


def _partition_bullets_by_block_counts(
    bullets: list[str],
    counts: list[int],
) -> list[list[str]]:
    """Split bullets across roles using ATS minimums — never cap to source resume counts."""
    return partition_experience_bullets_by_role(bullets, counts)


def _primary_role_header_index(para_block: list[int], paragraphs: list[Paragraph]) -> int | None:
    for idx in para_block:
        text = _paragraph_text(paragraphs[idx]).strip()
        if not text or _is_bullet_paragraph(paragraphs[idx]):
            continue
        if _is_experience_role_start(paragraphs[idx], text):
            return idx
    for idx in para_block:
        text = _paragraph_text(paragraphs[idx]).strip()
        if text and not _is_bullet_paragraph(paragraphs[idx]):
            return idx
    return None


def _apply_role_header_text(paragraph: Paragraph, text: str) -> None:
    """Experience role/company/date headers use full-line bold; no JD keyword bolding."""
    _replace_paragraph_text_plain(paragraph, text)
    _apply_role_header_bold(paragraph)


def _apply_role_header_bold(paragraph: Paragraph) -> None:
    """Keep role header lines bold in the exported Word file."""
    for run in paragraph.runs:
        if (run.text or "").strip():
            run.bold = True


def _paragraph_ppr(paragraph: Paragraph):
    from docx.oxml import OxmlElement

    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        paragraph._element.insert(0, p_pr)
    return p_pr


def _sync_paragraph_indent(source: Paragraph, target: Paragraph) -> None:
    """Align a company/meta line with the role title line above it (left edge)."""
    if source._element is target._element:
        return
    src_ppr = source._element.find(qn("w:pPr"))
    if src_ppr is None:
        return
    tgt_ppr = _paragraph_ppr(target)
    for tag in ("w:ind", "w:tabs", "w:jc"):
        qtag = qn(tag)
        for old in tgt_ppr.findall(qtag):
            tgt_ppr.remove(old)
        src_el = src_ppr.find(qtag)
        if src_el is not None:
            tgt_ppr.append(deepcopy(src_el))
    try:
        src_fmt = source.paragraph_format
        tgt_fmt = target.paragraph_format
        tgt_fmt.left_indent = src_fmt.left_indent
        tgt_fmt.right_indent = src_fmt.right_indent
        tgt_fmt.first_line_indent = src_fmt.first_line_indent
    except Exception:
        pass


def _company_line_index(
    para_block: list[int],
    paragraphs: list[Paragraph],
    header_idx: int,
) -> int | None:
    """First company/location line in a role block (non-bullet, not the role header)."""
    for idx in para_block:
        if idx == header_idx:
            continue
        para = paragraphs[idx]
        text = _paragraph_text(para).strip()
        if not text or _is_bullet_paragraph(para):
            continue
        if _is_experience_role_start(para, text):
            continue
        return idx
    return None


def _align_experience_company_indent(
    paragraphs: list[Paragraph],
    para_block: list[int],
    header_idx: int | None,
) -> None:
    if header_idx is None:
        return
    company_idx = _company_line_index(para_block, paragraphs, header_idx)
    if company_idx is None:
        return
    _sync_paragraph_indent(paragraphs[header_idx], paragraphs[company_idx])


def _align_cell_company_indent(cell, role_para: Paragraph | None = None) -> None:
    """Align the company line in a table cell with the role title paragraph."""
    role = role_para
    company_para: Paragraph | None = None
    for para in cell.paragraphs:
        text = _paragraph_text(para).strip()
        if not text or _is_bulletish_text(text):
            continue
        if role is None:
            role = para
            continue
        company_para = para
        break
    if role is not None and company_para is not None:
        _sync_paragraph_indent(role, company_para)


def _update_experience_bullets_only(
    doc: Document,
    indices: list[int],
    new_text: str,
    highlight_terms: list[str] | None = None,
    enable_bold: bool = True,
    bullets_per_role: list[int] | None = None,
) -> None:
    """
    Update bullet paragraphs in each work-experience role block only.
    Role/company/date headers are left unchanged from the source resume.
    """
    if not indices or not new_text.strip():
        return

    paragraphs = _all_document_paragraphs(doc)
    para_blocks = _group_experience_paragraph_indices(doc, indices)
    target_counts = _resolve_experience_bullet_targets(doc, indices, bullets_per_role)
    llm_bullets = _bullet_lines(new_text)
    bullets_per_block = _partition_bullets_by_block_counts(llm_bullets, target_counts)
    template_header: Paragraph | None = None

    for block_i, para_block in enumerate(para_blocks):
        header_idx = _primary_role_header_index(para_block, paragraphs)
        if header_idx is not None:
            header_para = paragraphs[header_idx]
            if template_header is None:
                template_header = header_para
            else:
                try:
                    header_para.style = template_header.style
                except Exception:
                    pass
            _apply_role_header_bold(header_para)

        _align_experience_company_indent(paragraphs, para_block, header_idx)

        bullet_indices = [idx for idx in para_block if _is_bullet_paragraph(paragraphs[idx])]
        if not bullet_indices:
            bullet_indices = [
                idx
                for idx in para_block
                if _BULLET_CHAR_RE.match(_paragraph_text(paragraphs[idx]).strip())
            ]
        if not bullet_indices:
            continue
        new_bullets = bullets_per_block[block_i] if block_i < len(bullets_per_block) else []
        slot_paragraphs = [paragraphs[idx] for idx in bullet_indices]
        template_para = slot_paragraphs[0]
        _apply_bullet_updates(
            slot_paragraphs,
            new_bullets,
            template_para,
            highlight_terms=highlight_terms,
            enable_bold=enable_bold,
        )


def _apply_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    highlight_terms: list[str] | None = None,
    plain: bool = False,
    selective_highlight: bool = False,
    bold_metrics: bool = False,
    enable_bold: bool = True,
) -> None:
    if plain or not enable_bold:
        _replace_paragraph_text_plain(paragraph, text)
    elif highlight_terms:
        _replace_paragraph_text_with_highlights(
            paragraph,
            text,
            highlight_terms,
            auto_tech_and_metrics=not selective_highlight,
            auto_metrics=bold_metrics,
        )
    else:
        _replace_paragraph_text_inplace(paragraph, text)


def _update_lines_index_preserving(
    doc: Document,
    indices: list[int],
    new_text: str,
    *,
    highlight_terms: list[str] | None = None,
    plain: bool = False,
    selective_highlight: bool = False,
    enable_bold: bool = True,
    clear_unused: bool = False,
) -> None:
    """
    Update paragraph i with line i only — preserves structure, styles, and extra paragraphs.
    When clear_unused is True, blank out skill paragraphs beyond the new line count.
    """
    lines = _split_tailored_lines(new_text)
    if not lines or not indices:
        return
    paragraphs = _all_document_paragraphs(doc)
    if len(indices) == 1 and len(lines) > 1:
        _apply_paragraph_text(
            paragraphs[indices[0]],
            " ".join(lines),
            highlight_terms=highlight_terms,
            plain=plain,
            selective_highlight=selective_highlight,
            enable_bold=enable_bold,
        )
        if clear_unused:
            for idx in indices[1:]:
                if idx < len(paragraphs):
                    _clear_paragraph_visible(paragraphs[idx])
        return
    for slot_i, idx in enumerate(indices):
        if idx >= len(paragraphs):
            break
        para = paragraphs[idx]
        if slot_i < len(lines):
            line = lines[slot_i]
            if _is_bullet_paragraph(para):
                line = _format_bullet_for_paragraph(para, line)
            else:
                existing = _paragraph_text(para).strip()
                if existing and _BULLET_CHAR_RE.match(existing) and not _BULLET_CHAR_RE.match(line):
                    line = _format_bullet_for_paragraph(para, line)
            display = _display_line_for_paragraph(line, para) if not _is_bullet_paragraph(para) else line
            _apply_paragraph_text(
                para,
                display,
                highlight_terms=highlight_terms,
                plain=plain,
                selective_highlight=selective_highlight,
                enable_bold=enable_bold,
            )
        elif clear_unused:
            if _paragraph_text(para).strip():
                _clear_paragraph_visible(para)


def _normalize_section_body_indices(
    body_indices: dict[str, list[int]],
    section_header_indices: dict[str, int],
    contact_paragraph_indices: list[int],
) -> None:
    """Keep each section's editable paragraphs strictly inside its header boundaries."""
    contact_set = set(contact_paragraph_indices)
    ordered_sections = (
        "professional_summary",
        "skills",
        "professional_experience",
        "education",
        "other",
    )
    header_positions = {
        name: section_header_indices[name]
        for name in ordered_sections
        if name in section_header_indices
    }

    for section in ordered_sections:
        indices = body_indices.get(section, [])
        if not indices:
            continue
        start = header_positions.get(section)
        if start is not None:
            indices = [i for i in indices if i > start]
        later_headers = [
            pos for name, pos in header_positions.items() if name != section and pos > (start or -1)
        ]
        if later_headers:
            end = min(later_headers)
            indices = [i for i in indices if i < end]
        if section != "contact":
            indices = [i for i in indices if i not in contact_set]
        body_indices[section] = indices


def _reassign_trailing_education_indices(
    body_indices: dict[str, list[int]],
    parsed: ParsedResume,
) -> None:
    if body_indices.get("education"):
        return
    edu_count = len(_split_tailored_lines(parsed.education))
    if edu_count <= 0:
        return
    skill_idx = body_indices.get("skills", [])
    if len(skill_idx) < edu_count:
        return
    body_indices["education"] = skill_idx[-edu_count:]
    body_indices["skills"] = skill_idx[:-edu_count]


def _normalize_table_label(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _table_row_label(table, row_idx: int = 0) -> str:
    if row_idx >= len(table.rows):
        return ""
    parts: list[str] = []
    seen: set[str] = set()
    for cell in table.rows[row_idx].cells:
        label = _normalize_table_label(cell.text)
        if label and label not in seen:
            seen.add(label)
            parts.append(label)
    return " | ".join(parts)


def _is_bulletish_text(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return bool(_BULLET_CHAR_RE.match(stripped))


def _looks_like_skill_category_line(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    if _SKILL_CATEGORY_RE.match(stripped):
        return True
    return _is_bulletish_text(stripped) and ":" in stripped.split("\n", 1)[0][:60]


def _cell_bullet_paragraph_indices(
    cell,
    *,
    after_para_idx: int | None = None,
    before_para_idx: int | None = None,
) -> list[int]:
    """Bullet paragraphs in a table cell (● markers, list styles, or dash prefixes)."""
    indices: list[int] = []
    for idx, para in enumerate(cell.paragraphs):
        if after_para_idx is not None and idx <= after_para_idx:
            continue
        if before_para_idx is not None and idx >= before_para_idx:
            continue
        if _is_bullet_paragraph(para) or _is_bulletish_text(_paragraph_text(para)):
            indices.append(idx)
    return indices


def _cell_role_header_paragraph_indices(cell) -> list[int]:
    """Non-bullet header paragraphs inside a table cell that start a new job (supports 3+ roles per cell)."""
    headers: list[int] = []
    seen_bullets_since_last_header = False
    for idx, para in enumerate(cell.paragraphs):
        text = _paragraph_text(para).strip()
        if not text:
            continue
        if _is_bulletish_text(text):
            seen_bullets_since_last_header = True
            continue
        is_header = is_experience_role_header_line(text) or _looks_like_job_title_line(text)
        if not is_header:
            continue
        if not headers:
            headers.append(idx)
            seen_bullets_since_last_header = False
            continue
        if seen_bullets_since_last_header or is_experience_role_header_line(text):
            headers.append(idx)
            seen_bullets_since_last_header = False
    return headers


def _expand_experience_row_refs(doc: Document, rows: list[ExperienceRowRef]) -> list[ExperienceRowRef]:
    """Split table rows that contain multiple jobs in one cell into separate role refs."""
    expanded: list[ExperienceRowRef] = []
    for ref in rows:
        cell = doc.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]]
        header_indices = _cell_role_header_paragraph_indices(cell)
        if len(header_indices) <= 1:
            expanded.append(ref)
            continue
        for header_idx in header_indices:
            expanded.append(
                ExperienceRowRef(
                    table_idx=ref.table_idx,
                    row_idx=ref.row_idx,
                    content_cols=ref.content_cols,
                    role_header_para_idx=header_idx,
                )
            )
    return expanded


def _detect_bullet_prefix(cell) -> str:
    for para in cell.paragraphs:
        text = _paragraph_text(para).strip()
        if not text:
            continue
        match = _BULLET_CHAR_RE.match(text)
        if match:
            return match.group(0).rstrip()
    return "•"


def _format_bullet_line(prefix: str, line: str) -> str:
    body = _BULLET_PREFIX_RE.sub("", line).strip()
    body = re.sub(r"^[\-•*–—\u2022]\s*", "", body).strip()
    sep = "" if prefix.endswith(" ") else " "
    return f"{prefix}{sep}{body}".strip()


def _find_work_experience_table(doc: Document) -> tuple[int, int | None] | None:
    for ti, table in enumerate(doc.tables):
        label = _table_row_label(table, 0)
        if not _TABLE_WORK_EXP_RE.search(label):
            continue
        edu_row: int | None = None
        for ri in range(1, len(table.rows)):
            row_label = _table_row_label(table, ri)
            first = row_label.split("|", 1)[0].strip()
            if _TABLE_EDUCATION_RE.match(first):
                edu_row = ri
                break
        return ti, edu_row
    return None


def _detect_table_layout(doc: Document) -> tuple[list[int], list[int], list[ExperienceRowRef]] | None:
    """Detect summary/skills paragraphs and experience table rows (table-based templates)."""
    work = _find_work_experience_table(doc)
    if work is None:
        return None
    table_idx, edu_row = work

    summary_indices: list[int] = []
    skills_indices: list[int] = []
    skill_start: int | None = None

    for idx, para in enumerate(doc.paragraphs):
        text = _paragraph_text(para).strip()
        if not text:
            continue
        if _looks_like_skill_category_line(text):
            if skill_start is None:
                skill_start = idx
            skills_indices.append(idx)
        elif skill_start is None and len(text) > 60 and not _is_bulletish_text(text):
            summary_indices.append(idx)

    if not summary_indices:
        for idx, para in enumerate(doc.paragraphs):
            if _paragraph_text(para).strip():
                summary_indices = [idx]
                break

    exp_table = doc.tables[table_idx]
    end_row = edu_row if edu_row is not None else len(exp_table.rows)
    exp_rows: list[ExperienceRowRef] = []
    for ri in range(1, end_row):
        row_label = _table_row_label(exp_table, ri)
        first = row_label.split("|", 1)[0].strip()
        if _TABLE_EDUCATION_RE.match(first):
            break
        row = exp_table.rows[ri]
        content_cols = _unique_content_col_indices(row)
        content_cell = row.cells[content_cols[0]]
        if not content_cell.text.strip():
            continue
        exp_rows.append(ExperienceRowRef(table_idx=table_idx, row_idx=ri, content_cols=content_cols))

    exp_rows = _expand_experience_row_refs(doc, exp_rows)

    if not exp_rows and not summary_indices and not skills_indices:
        return None
    return summary_indices, skills_indices, exp_rows


def _experience_text_from_table_row(doc: Document, ref: ExperienceRowRef) -> list[str]:
    cell = doc.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]]
    header_indices = _cell_role_header_paragraph_indices(cell)
    start_idx = 0
    end_idx = len(cell.paragraphs)
    if ref.role_header_para_idx is not None:
        start_idx = ref.role_header_para_idx
        for header_idx in header_indices:
            if header_idx > ref.role_header_para_idx:
                end_idx = header_idx
                break
    lines: list[str] = []
    header_done = False
    for para_idx, para in enumerate(cell.paragraphs):
        if para_idx < start_idx or para_idx >= end_idx:
            continue
        text = _paragraph_text(para).strip()
        if not text:
            continue
        if _is_bulletish_text(text):
            header_done = True
            lines.append(_format_bullet_line("- ", text))
        elif not header_done:
            lines.append(text)
    date_cell = doc.tables[ref.table_idx].rows[ref.row_idx].cells[-1].text.strip()
    if date_cell and len(header_indices) <= 1:
        lines.append(date_cell)
    return lines


def _parsed_from_table_layout(
    doc: Document,
    summary_indices: list[int],
    skills_indices: list[int],
    exp_rows: list[ExperienceRowRef],
) -> ParsedResume:
    contact_lines: list[str] = []
    if doc.tables:
        for row in doc.tables[0].rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                contact_lines.append(" | ".join(dict.fromkeys(cells)))

    summary = "\n".join(_paragraph_text(doc.paragraphs[i]) for i in summary_indices).strip()
    skills = "\n".join(_paragraph_text(doc.paragraphs[i]) for i in skills_indices).strip()

    exp_blocks: list[str] = []
    for ref in exp_rows:
        block_lines = _experience_text_from_table_row(doc, ref)
        if block_lines:
            exp_blocks.append("\n".join(block_lines))
    experience = "\n\n".join(exp_blocks).strip()

    education_lines: list[str] = []
    other_lines: list[str] = []
    work = _find_work_experience_table(doc)
    if work:
        table_idx, edu_row = work
        if edu_row is not None:
            exp_table = doc.tables[table_idx]
            for ri in range(edu_row + 1, len(exp_table.rows)):
                row_text = " | ".join(
                    c.text.strip() for c in exp_table.rows[ri].cells if c.text.strip()
                )
                if row_text:
                    education_lines.append(row_text)

    for idx, para in enumerate(doc.paragraphs):
        text = _paragraph_text(para).strip()
        if not text or idx in summary_indices or idx in skills_indices:
            continue
        if _looks_like_skill_category_line(text):
            continue
        if idx in summary_indices:
            continue
        other_lines.append(text)

    return ParsedResume(
        contact="\n".join(contact_lines).strip(),
        professional_summary=summary,
        professional_experience=experience,
        skills=skills,
        education="\n".join(education_lines).strip(),
        other="\n".join(other_lines).strip(),
    )


def _apply_table_layout_to_parse(
    doc: Document,
    body_indices: dict[str, list[int]],
    section_header_indices: dict[str, int],
) -> tuple[ParsedResume, list[ExperienceRowRef], str]:
    layout = _detect_table_layout(doc)
    if layout is None:
        raise ValueError("no table layout")
    summary_indices, skills_indices, exp_rows = layout

    body_indices["professional_summary"] = summary_indices
    body_indices["skills"] = skills_indices
    body_indices["professional_experience"] = []
    body_indices["education"] = []
    body_indices["other"] = []

    for ti, table in enumerate(doc.tables):
        label = _table_row_label(table, 0)
        if _TABLE_SUMMARY_RE.match(label.split("|", 1)[0].strip()):
            section_header_indices.setdefault("professional_summary", -ti - 1)
        elif _TABLE_SKILLS_RE.match(label.split("|", 1)[0].strip()):
            section_header_indices.setdefault("skills", -ti - 1)
        elif _TABLE_WORK_EXP_RE.search(label):
            section_header_indices.setdefault("professional_experience", -ti - 1)

    parsed = _parsed_from_table_layout(doc, summary_indices, skills_indices, exp_rows)
    plain_parts = [parsed.contact, parsed.professional_summary, parsed.skills, parsed.professional_experience]
    plain_text = "\n\n".join(p for p in plain_parts if p).strip()
    return parsed, exp_rows, plain_text


def parse_role_titles_from_table_rows(doc: Document, rows: list[ExperienceRowRef]) -> list[str]:
    """One role title per experience table row — first non-bullet line in the content cell."""
    titles: list[str] = []
    for ref in rows:
        cell = doc.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]]
        for para in cell.paragraphs:
            text = _paragraph_text(para).strip()
            if not text or _is_bulletish_text(text):
                continue
            if "|" in text:
                titles.append(text.split("|", 1)[0].strip())
            else:
                titles.append(text)
            break
    return titles


def _update_experience_table_rows(
    doc: Document,
    rows: list[ExperienceRowRef],
    new_text: str,
    highlight_terms: list[str] | None = None,
    enable_bold: bool = True,
    bullets_per_role: list[int] | None = None,
) -> None:
    """Update bullet paragraphs inside work-experience table rows only."""
    if not rows or not new_text.strip():
        return

    role_count = len(rows)
    template_slots = [
        len(
            _cell_bullet_paragraph_indices(
                doc.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]],
                after_para_idx=ref.role_header_para_idx,
            )
        )
        for ref in rows
    ]
    ats_targets = (
        bullets_per_role
        if bullets_per_role and len(bullets_per_role) == role_count
        else default_ats_bullets_per_role(role_count)
    )
    target_counts = effective_bullets_per_role(ats_targets, template_slots)
    llm_bullets = _bullet_lines(new_text)
    bullets_per_row = _partition_bullets_by_block_counts(llm_bullets, target_counts)

    for row_i, ref in enumerate(rows):
        cell = doc.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]]
        role_para: Paragraph | None = None
        header_indices = _cell_role_header_paragraph_indices(cell)
        next_header_idx: int | None = None
        if ref.role_header_para_idx is not None:
            for hi, header_idx in enumerate(header_indices):
                if header_idx == ref.role_header_para_idx:
                    role_para = cell.paragraphs[header_idx]
                    if hi + 1 < len(header_indices):
                        next_header_idx = header_indices[hi + 1]
                    break
        if role_para is None:
            for para in cell.paragraphs:
                text = _paragraph_text(para).strip()
                if text and not _is_bulletish_text(text):
                    role_para = para
                    break
        if role_para is not None:
            _apply_role_header_bold(role_para)
        _align_cell_company_indent(cell, role_para)

        bullet_indices = _cell_bullet_paragraph_indices(
            cell,
            after_para_idx=ref.role_header_para_idx,
            before_para_idx=next_header_idx,
        )
        if not bullet_indices:
            continue
        new_bullets = bullets_per_row[row_i] if row_i < len(bullets_per_row) else []
        slot_paragraphs = [cell.paragraphs[para_idx] for para_idx in bullet_indices]
        template_para = slot_paragraphs[0]
        _apply_bullet_updates(
            slot_paragraphs,
            new_bullets,
            template_para,
            highlight_terms=highlight_terms,
            enable_bold=enable_bold,
        )

        for col in ref.content_cols[1:]:
            row_cells = doc.tables[ref.table_idx].rows[ref.row_idx].cells
            if col < len(row_cells):
                _sync_cell_paragraphs(row_cells[col], cell)


def _unique_content_col_indices(row) -> tuple[int, ...]:
    """Column indices for editable content, deduped when Word merges cells (same tc repeated)."""
    cols: list[int] = []
    seen_tc: set[int] = set()
    for ci, cell in enumerate(row.cells):
        tc_id = id(cell._tc)
        if tc_id in seen_tc:
            continue
        seen_tc.add(tc_id)
        cols.append(ci)
    return tuple(cols[:1] if cols else (0,))


def _sync_cell_paragraphs(dup_cell, source_cell) -> None:
    """Mirror paragraph runs (including bold highlights) into duplicate merged cells."""
    if dup_cell._tc is source_cell._tc:
        return
    src_paras = source_cell.paragraphs
    dup_paras = dup_cell.paragraphs
    for i, src_para in enumerate(src_paras):
        if i < len(dup_paras):
            _sync_paragraph_runs(dup_paras[i], src_para)


def _has_editable_docx_targets(
    section_body_indices: dict[str, list[int]],
    experience_table_rows: list[ExperienceRowRef],
) -> bool:
    if experience_table_rows:
        return True
    return any(section_body_indices.get(section) for section in _EDITABLE_DOCX_SECTIONS)


def _field_hyperlink_urls_from_xml(xml_bytes: bytes) -> list[str]:
    """Extract URLs stored as Word HYPERLINK field codes instead of relationships."""
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return []
    instructions = " ".join(
        (element.text or "").strip()
        for element in root.iter()
        if element.tag.rsplit("}", 1)[-1] == "instrText" and (element.text or "").strip()
    )
    urls: list[str] = []
    for match in re.finditer(r'\bHYPERLINK\s+(?:"([^"]+)"|(\S+))', instructions, re.I):
        url = (match.group(1) or match.group(2) or "").strip()
        if url.lower().startswith(("http://", "https://")) and url not in urls:
            urls.append(url)
    return urls


def extract_http_links_from_docx(docx_bytes: bytes) -> list[tuple[str, str]]:
    if not docx_bytes:
        return []
    labeled: list[tuple[str, str]] = []
    seen: set[str] = set()
    try:
        doc = Document(BytesIO(docx_bytes))
    except (OSError, ValueError, BadZipFile):
        doc = None
    if doc is not None:
        for paragraph in doc.paragraphs:
            for hyperlink in paragraph._element.xpath(".//w:hyperlink"):
                r_id = hyperlink.get(qn("r:id"))
                if not r_id:
                    continue
                try:
                    url = paragraph.part.rels[r_id].target_ref.strip()
                except KeyError:
                    continue
                if not url.lower().startswith(("http://", "https://")):
                    continue
                key = url.lower().rstrip("/")
                if key in seen:
                    continue
                seen.add(key)
                label = "LinkedIn" if "linkedin.com" in key else "GitHub" if "github.com" in key else "Link"
                labeled.append((label, url))
    # Hyperlinks may live in header/footer parts, tables, shapes, or text boxes and
    # therefore never appear in ``doc.paragraphs``. Read every external relationship.
    try:
        with ZipFile(BytesIO(docx_bytes)) as archive:
            for name in archive.namelist():
                if not name.startswith("word/") or not name.endswith(".rels"):
                    continue
                try:
                    root = ET.fromstring(archive.read(name))
                except ET.ParseError:
                    continue
                for relationship in root:
                    url = (relationship.attrib.get("Target") or "").strip()
                    if (relationship.attrib.get("TargetMode") or "").lower() != "external":
                        continue
                    if not url.lower().startswith(("http://", "https://")):
                        continue
                    key = url.lower().rstrip("/")
                    if key in seen:
                        continue
                    seen.add(key)
                    labeled.append((_link_label_for("", url), url))
            # Some Word generators use field codes such as
            # HYPERLINK "https://linkedin.com/in/name" with no .rels entry.
            for name in archive.namelist():
                if not name.startswith("word/") or not name.endswith(".xml"):
                    continue
                for url in _field_hyperlink_urls_from_xml(archive.read(name)):
                    key = url.lower().rstrip("/")
                    if key in seen:
                        continue
                    seen.add(key)
                    labeled.append((_link_label_for("", url), url))
    except (OSError, ValueError, BadZipFile):
        pass
    return labeled


def parse_resume_from_docx(data: bytes) -> DocxResumeDocument:
    if not _looks_like_docx(data):
        raise ValueError("Not a valid .docx file (expected a ZIP-based Word document).")

    doc = Document(BytesIO(data))
    paragraphs = _all_document_paragraphs(doc)
    buckets: dict[str, list[str]] = {
        "contact": [],
        "professional_summary": [],
        "professional_experience": [],
        "skills": [],
        "education": [],
        "other": [],
    }
    body_indices: dict[str, list[int]] = {key: [] for key in buckets}
    section_header_indices: dict[str, int] = {}
    contact_paragraph_indices: list[int] = []
    current = "contact"
    plain_lines: list[str] = []
    contact_ready = False

    for idx, paragraph in enumerate(paragraphs):
        line = _line_from_paragraph(paragraph)
        if line:
            stripped = line.strip()
            sec = match_section_header(stripped, in_contact=(current == "contact"))
            if sec:
                current = sec
                section_header_indices[sec] = idx
                if is_pure_section_header(stripped):
                    if sec == "other":
                        buckets[current].append(line)
                        body_indices[current].append(idx)
                        plain_lines.append(line)
                    continue
                if sec == "other":
                    buckets[current].append(line)
                    body_indices[current].append(idx)
                    plain_lines.append(line)
                continue
            if current == "contact":
                if line_is_factual_contact(stripped) or _detect_contact_header_role_line(stripped):
                    buckets["contact"].append(line)
                    body_indices["contact"].append(idx)
                    contact_paragraph_indices.append(idx)
                    plain_lines.append(line)
                    contact_ready = True
                    continue
                implicit = implicit_section_after_contact(stripped)
                if implicit is None and _is_experience_role_start(paragraph, stripped):
                    implicit = "professional_experience"
                if implicit is None and _is_bullet_paragraph(paragraph):
                    if not line_is_factual_contact(stripped):
                        implicit = "professional_experience"
                if implicit:
                    current = implicit
                elif line_is_factual_contact(stripped) and not _looks_like_role_header_line(stripped):
                    contact_ready = True
                elif not contact_ready and len(stripped) < 80 and not _is_bullet_paragraph(paragraph):
                    contact_ready = True
            elif current == "professional_summary":
                implicit = implicit_section_after_contact(stripped)
                if implicit is None and _is_experience_role_start(paragraph, stripped):
                    implicit = "professional_experience"
                if implicit is None and _is_bullet_paragraph(paragraph):
                    implicit = "professional_experience"
                if implicit == "professional_experience":
                    current = implicit
            elif current == "professional_experience":
                implicit = implicit_section_after_contact(stripped)
                if implicit in ("skills", "education"):
                    current = implicit
            elif current == "education":
                implicit = implicit_section_after_contact(stripped)
                if implicit == "skills":
                    current = implicit
            buckets[current].append(line)
            body_indices[current].append(idx)
            plain_lines.append(line)
            if current == "contact":
                contact_paragraph_indices.append(idx)

    for line in _table_lines(doc):
        buckets[current].append(line)
        plain_lines.append(line)

    def join_bucket(key: str) -> str:
        return "\n".join(buckets[key]).strip()

    contact = join_bucket("contact")
    summary = join_bucket("professional_summary")
    experience = join_bucket("professional_experience")
    skills = join_bucket("skills")
    education = join_bucket("education")
    other = join_bucket("other")

    if not any([summary, experience, skills, education, other]) and contact:
        experience = contact
        contact = ""
        body_indices["professional_experience"] = list(body_indices["contact"])
        body_indices["contact"] = []
        contact_paragraph_indices = []

    education, other = _partition_education_and_other(education, other)
    education, experience = _separate_misplaced_jobs_from_education(education, experience)

    parsed = ParsedResume(
        contact=contact,
        professional_summary=summary,
        professional_experience=experience,
        skills=skills,
        education=education,
        other=other,
    )
    plain_text = "\n".join(plain_lines).strip()

    if not plain_text.strip():
        raise ValueError("Could not read any text from this Word document.")

    if not any([summary, experience, skills]) and plain_text:
        parsed = parse_resume_sections(plain_text)

    section_body_indices = {key: list(body_indices[key]) for key in _TAILOR_HEADER_SECTIONS}
    _normalize_section_body_indices(
        section_body_indices,
        section_header_indices,
        contact_paragraph_indices,
    )
    _reassign_trailing_education_indices(section_body_indices, parsed)

    experience_table_rows: list[ExperienceRowRef] = []
    table_layout = _detect_table_layout(doc)
    if table_layout is not None:
        try:
            parsed, experience_table_rows, plain_text = _apply_table_layout_to_parse(
                doc,
                section_body_indices,
                section_header_indices,
            )
        except ValueError:
            experience_table_rows = []

    text_role_count = len(split_experience_line_blocks(parsed.professional_experience))
    exp_indices = section_body_indices.get("professional_experience") or []
    grouped_role_count = len(_group_experience_paragraph_indices(doc, exp_indices)) if exp_indices else 0
    table_role_count = len(experience_table_rows)
    if grouped_role_count > 0:
        detected_role_count = max(grouped_role_count, table_role_count)
    else:
        detected_role_count = max(text_role_count, table_role_count) or None

    experience_bullet_slots: list[int] = []
    if experience_table_rows:
        for ref in experience_table_rows:
            cell = doc.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]]
            experience_bullet_slots.append(
                len(
                    _cell_bullet_paragraph_indices(
                        cell,
                        after_para_idx=ref.role_header_para_idx,
                    )
                )
            )
    elif exp_indices:
        experience_bullet_slots = _experience_template_bullet_slots(doc, exp_indices)
    if not experience_bullet_slots:
        experience_bullet_slots = [
            sum(1 for line in block if _BULLET_PREFIX_RE.match(line.strip()))
            for block in split_experience_line_blocks(parsed.professional_experience)
        ]

    return DocxResumeDocument(
        parsed=parsed,
        plain_text=plain_text,
        section_header_indices=section_header_indices,
        section_body_indices=section_body_indices,
        contact_paragraph_indices=contact_paragraph_indices,
        experience_table_rows=experience_table_rows,
        detected_role_count=detected_role_count,
        experience_bullet_slots=experience_bullet_slots,
    )


def _update_contact_inplace(
    doc: Document,
    contact_indices: list[int],
    new_text: str,
) -> None:
    paragraphs = _all_document_paragraphs(doc)
    valid_indices = [i for i in contact_indices if i < len(paragraphs)]
    if not valid_indices:
        return

    new_lines = [line for line in _split_tailored_lines(new_text) if not line_is_factual_contact(line)]
    header_indices = [
        i for i in valid_indices if not line_is_factual_contact(_paragraph_text(paragraphs[i]))
    ]
    new_i = 0
    for idx in header_indices:
        if new_i >= len(new_lines):
            break
        para = paragraphs[idx]
        _replace_paragraph_text_inplace(
            para,
            _display_line_for_paragraph(new_lines[new_i], para),
        )
        new_i += 1


def _update_summary_preserving_spacing(
    doc: Document,
    indices: list[int],
    new_text: str,
    *,
    highlight_terms: list[str] | None = None,
    enable_bold: bool = True,
) -> None:
    """
    Update summary text without clearing extra paragraph slots or changing paragraph spacing.
    Writes the full summary into the first slot only; leaves other summary paragraphs untouched.
    """
    lines = _split_tailored_lines(new_text)
    if not lines or not indices:
        return
    paragraphs = _all_document_paragraphs(doc)
    first_idx = indices[0]
    if first_idx >= len(paragraphs):
        return
    summary_text = " ".join(lines)
    _apply_paragraph_text(
        paragraphs[first_idx],
        summary_text,
        highlight_terms=highlight_terms,
        enable_bold=enable_bold,
    )


def _update_section_inplace(
    doc: Document,
    section_name: str,
    indices: list[int],
    new_text: str,
    source_text: str,
    highlight_terms: list[str] | None = None,
    enable_bold: bool = True,
) -> None:
    if section_name in _FROZEN_DOCX_SECTIONS:
        return
    if section_name not in _EDITABLE_DOCX_SECTIONS:
        return
    if section_name == "professional_experience":
        _update_experience_bullets_only(
            doc, indices, new_text, highlight_terms, enable_bold=enable_bold
        )
    elif section_name == "skills":
        _update_skills_preserving_template(
            doc,
            indices,
            new_text,
            source_text,
            highlight_terms=highlight_terms,
            enable_bold=enable_bold,
        )
    elif section_name == "professional_summary":
        _update_summary_preserving_spacing(
            doc,
            indices,
            new_text,
            highlight_terms=highlight_terms,
            enable_bold=enable_bold,
        )
    else:
        _update_lines_index_preserving(doc, indices, new_text)


def _ensure_tailored_experience_has_bullets(tailored_experience: str, source_experience: str) -> str:
    """Keep source bullet lines when the LLM returns headers only (prevents empty experience rows)."""
    tailored = (tailored_experience or "").strip()
    source = (source_experience or "").strip()
    if not tailored or not source:
        return tailored
    if _bullet_lines(tailored):
        return tailored
    source_bullets = _bullet_lines(source)
    if not source_bullets:
        return tailored
    return merge_experience_headers_with_bullets(source, tailored)


def _count_body_level_paragraphs(document_xml: str) -> int:
    children = _split_document_body_children(document_xml)
    if not children:
        return 0
    return sum(1 for child in children if child.startswith("<w:p"))


def _count_tables_in_document_xml(document_xml: str) -> int:
    children = _split_document_body_children(document_xml)
    if not children:
        return 0
    return sum(1 for child in children if child.startswith("<w:tbl"))


def _build_editable_body_para_xml_map(
    memory_doc: Document,
    editable_paragraph_indices: set[int],
) -> dict[int, str]:
    """Body-level paragraph XML from the in-memory doc (summary/skills — not inside tables)."""
    editable: dict[int, str] = {}
    paragraphs = memory_doc.paragraphs
    for idx in editable_paragraph_indices:
        if 0 <= idx < len(paragraphs):
            editable[idx] = paragraphs[idx]._element.xml
    return editable


def _rebuild_orig_document_selective(
    orig_document_xml: str,
    memory_doc: Document,
    *,
    editable_paragraph_indices: set[int],
    editable_table_indices: set[int],
) -> str | None:
    """
    Rebuild document.xml from the original upload layout, swapping only edited
    body paragraphs and tables from the in-memory doc. Never adopts a flattened save().
    """
    orig_children = _split_document_body_children(orig_document_xml)
    if not orig_children:
        return None

    orig_body_para_count = _count_body_level_paragraphs(orig_document_xml)
    if len(memory_doc.paragraphs) != orig_body_para_count:
        return None

    editable_para_xml = _build_editable_body_para_xml_map(memory_doc, editable_paragraph_indices)
    mem_tables = list(memory_doc.tables)

    body_para_idx = 0
    tbl_idx = 0
    merged_parts: list[str] = []

    for child in orig_children:
        if child.startswith("<w:sectPr"):
            merged_parts.append(child)
            continue
        if child.startswith("<w:tbl"):
            if tbl_idx in editable_table_indices and tbl_idx < len(mem_tables):
                merged_parts.append(
                    _merge_table_xml_preserving_ppr(child, mem_tables[tbl_idx]._element.xml)
                )
            else:
                merged_parts.append(child)
            tbl_idx += 1
            continue
        if child.startswith("<w:p"):
            mod_xml = editable_para_xml.get(body_para_idx)
            if mod_xml is not None:
                merged_parts.append(_merge_paragraph_xml_preserving_ppr(child, mod_xml))
            else:
                merged_parts.append(child)
            body_para_idx += 1
            continue
        merged_parts.append(child)

    body_match = re.search(r"(<w:body>).*?(</w:body>)", orig_document_xml, re.S)
    if not body_match:
        return None
    merged_body = body_match.group(1) + "".join(merged_parts) + body_match.group(2)
    result = orig_document_xml[: body_match.start()] + merged_body + orig_document_xml[body_match.end() :]
    return _postprocess_document_xml(result)


def _first_table_xml(document_xml: str) -> str | None:
    body_match = re.search(r"<w:body>(.*)</w:body>", document_xml, re.S)
    if not body_match:
        return None
    inner = body_match.group(1)
    pos = inner.find("<w:tbl")
    if pos < 0:
        return None
    end = _find_w_element_end(inner, pos, "tbl")
    return inner[pos:end] if end > pos else None


def _replace_first_table_xml(document_xml: str, header_table_xml: str) -> str:
    body_match = re.search(r"(<w:body>)(.*?)(</w:body>)", document_xml, re.S)
    if not body_match:
        return document_xml
    inner = body_match.group(2)
    pos = inner.find("<w:tbl")
    if pos < 0:
        return document_xml
    end = _find_w_element_end(inner, pos, "tbl")
    if end < 0:
        return document_xml
    new_inner = inner[:pos] + header_table_xml + inner[end:]
    return (
        document_xml[: body_match.start()]
        + body_match.group(1)
        + new_inner
        + body_match.group(3)
        + document_xml[body_match.end() :]
    )


def _find_w_element_end(xml: str, start: int, local_name: str) -> int:
    """Return the end index of a <w:local_name ...>...</w:local_name> element."""
    open_prefix = f"<w:{local_name}"
    close_tag = f"</w:{local_name}>"
    if not xml.startswith(open_prefix, start):
        return -1
    after = start + len(open_prefix)
    if after >= len(xml) or xml[after] not in (">", " ", "/"):
        return -1

    pos = after
    depth = 1
    while pos < len(xml) and depth > 0:
        next_open = _find_next_w_open(xml, pos, local_name)
        next_close = xml.find(close_tag, pos)
        if next_close < 0:
            return -1
        if next_open >= 0 and next_open < next_close:
            depth += 1
            pos = next_open + len(open_prefix)
            continue
        depth -= 1
        pos = next_close + len(close_tag)
    return pos


def _find_next_w_open(xml: str, pos: int, local_name: str) -> int:
    open_prefix = f"<w:{local_name}"
    idx = pos
    while True:
        idx = xml.find(open_prefix, idx)
        if idx < 0:
            return -1
        after = idx + len(open_prefix)
        if after < len(xml) and xml[after] in (">", " ", "/"):
            return idx
        idx = after


def _split_document_body_children(document_xml: str) -> list[str] | None:
    body_match = re.search(r"<w:body>(.*)</w:body>", document_xml, re.S)
    if not body_match:
        return None

    inner = body_match.group(1)
    children: list[str] = []
    pos = 0
    length = len(inner)
    while pos < length:
        while pos < length and inner[pos].isspace():
            pos += 1
        if pos >= length:
            break

        if inner.startswith("<w:p", pos):
            end = _find_w_element_end(inner, pos, "p")
        elif inner.startswith("<w:tbl", pos):
            end = _find_w_element_end(inner, pos, "tbl")
        elif inner.startswith("<w:sectPr", pos):
            end = _find_w_element_end(inner, pos, "sectPr")
        else:
            return None

        if end < 0:
            return None
        children.append(inner[pos:end])
        pos = end

    return children or None


def _paragraph_plain_text_from_xml(paragraph_xml: str) -> str:
    parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", paragraph_xml, re.S)
    return "".join(parts).strip()


def _minimize_empty_paragraph_xml(paragraph_xml: str) -> str:
    """Collapse empty paragraphs to near-zero height in merged document.xml."""
    if not paragraph_xml.startswith("<w:p"):
        return paragraph_xml
    xml = paragraph_xml
    spacing_el = '<w:spacing w:before="0" w:after="0" w:line="1" w:lineRule="exact"/>'
    if "<w:pPr" in xml:
        xml = re.sub(r"<w:numPr\b[^>]*/>", "", xml)
        xml = re.sub(r"<w:numPr\b[^>]*>.*?</w:numPr>", "", xml, flags=re.S)
        if re.search(r"<w:spacing\b", xml):
            xml = re.sub(r"<w:spacing[^>]*/>", spacing_el, xml, count=1)
        else:
            xml = re.sub(r"(<w:pPr[^>]*>)", r"\1" + spacing_el, xml, count=1)
    else:
        xml = re.sub(r"(<w:p[^>]*>)", r"\1<w:pPr>" + spacing_el + "</w:pPr>", xml, count=1)

    def _vanish_run(run_xml: str) -> str:
        if "<w:rPr" in run_xml:
            if "<w:vanish" not in run_xml:
                run_xml = re.sub(r"(<w:rPr[^>]*>)", r"\1<w:vanish/>", run_xml, count=1)
            if re.search(r"<w:sz\b", run_xml):
                run_xml = re.sub(r"<w:sz\b[^>]*/>", '<w:sz w:val="2"/>', run_xml, count=1)
            else:
                run_xml = re.sub(r"(<w:rPr[^>]*>)", r'\1<w:sz w:val="2"/>', run_xml, count=1)
        else:
            run_xml = run_xml.replace(
                "<w:r>",
                '<w:r><w:rPr><w:vanish/><w:sz w:val="2"/></w:rPr>',
                1,
            )
        return run_xml

    if re.search(r"<w:r\b", xml):
        xml = re.sub(r"<w:r\b[^>]*>.*?</w:r>", lambda m: _vanish_run(m.group(0)), xml, flags=re.S)
    else:
        xml = xml.replace("</w:p>", '<w:r><w:rPr><w:vanish/><w:sz w:val="2"/></w:rPr><w:t></w:t></w:r></w:p>')
    return xml


def _collapse_spacing_in_paragraph_xml(paragraph_xml: str) -> str:
    """Alias for empty-paragraph minimization in XML merge passes."""
    return _minimize_empty_paragraph_xml(paragraph_xml)


def _strip_italic_from_paragraph_xml(paragraph_xml: str) -> str:
    """Remove italic and force normal emphasis on role-header paragraphs."""
    xml = re.sub(r"<w:i\b[^>]*/>", "", paragraph_xml)
    xml = re.sub(r"<w:iCs\b[^>]*/>", "", xml)
    xml = re.sub(r"<w:i\b[^>]*>.*?</w:i>", "", xml, flags=re.S)
    xml = re.sub(r"<w:iCs\b[^>]*>.*?</w:iCs>", "", xml, flags=re.S)

    def _force_run_non_italic(run_xml: str) -> str:
        if "<w:rPr" in run_xml:
            if re.search(r"<w:i\b", run_xml):
                run_xml = re.sub(r"<w:i\b[^>]*/>", '<w:i w:val="0"/>', run_xml)
                run_xml = re.sub(r"<w:iCs\b[^>]*/>", '<w:iCs w:val="0"/>', run_xml)
            else:
                run_xml = re.sub(
                    r"(<w:rPr[^>]*>)",
                    r'\1<w:i w:val="0"/><w:iCs w:val="0"/>',
                    run_xml,
                    count=1,
                )
        else:
            run_xml = run_xml.replace(
                "<w:r>",
                '<w:r><w:rPr><w:i w:val="0"/><w:iCs w:val="0"/></w:rPr>',
                1,
            )
        return run_xml

    return re.sub(r"<w:r\b[^>]*>.*?</w:r>", lambda m: _force_run_non_italic(m.group(0)), xml, flags=re.S)


def _paragraph_ppr_xml(paragraph_xml: str) -> str:
    match = re.search(r"<w:pPr\b[^>]*>.*?</w:pPr>", paragraph_xml, re.S)
    if match:
        return match.group(0)
    match = re.search(r"<w:pPr\b[^>]*/>", paragraph_xml)
    return match.group(0) if match else ""


def _paragraph_body_without_ppr(paragraph_xml: str) -> str:
    body = paragraph_xml
    body = re.sub(r"<w:pPr\b[^>]*>.*?</w:pPr>", "", body, count=1, flags=re.S)
    body = re.sub(r"<w:pPr\b[^>]*/>", "", body, count=1)
    body = re.sub(r"^<w:p[^>]*>", "", body)
    body = re.sub(r"</w:p>\s*$", "", body)
    return body


def _patch_paragraph_content_preserve_ppr(orig_xml: str, mod_xml: str) -> str:
    """Swap run content from mod while keeping original w:pPr (spacing/indents unchanged)."""
    close = orig_xml.rfind("</w:p>")
    open_end = orig_xml.find(">")
    if close < 0 or open_end < 0:
        return mod_xml
    ppr = _paragraph_ppr_xml(orig_xml)
    mod_body = _paragraph_body_without_ppr(mod_xml)
    if ppr:
        return orig_xml[: open_end + 1] + ppr + mod_body + orig_xml[close:]
    return mod_xml


def _merge_paragraph_xml_preserving_ppr(orig_xml: str, mod_xml: str) -> str:
    """
    Content-only merge: never replace layout/spacing from the upload.
    - Empty original → keep original (template spacer).
    - Cleared bullet slot → minimize height.
    - Updated text → new runs, original w:pPr.
    """
    orig_text = _paragraph_plain_text_from_xml(orig_xml)
    mod_text = _paragraph_plain_text_from_xml(mod_xml)
    if orig_text == mod_text:
        return orig_xml
    if not orig_text:
        return orig_xml
    if not mod_text:
        cleared = _paragraph_xml_with_empty_text(orig_xml)
        return _minimize_empty_paragraph_xml(cleared)
    return _patch_paragraph_content_preserve_ppr(orig_xml, mod_xml)


def _paragraph_xml_with_empty_text(paragraph_xml: str) -> str:
    """Clear visible text but keep original w:pPr spacing/indents (unused bullet slots)."""
    if not paragraph_xml.startswith("<w:p"):
        return paragraph_xml
    xml = re.sub(
        r"(<w:t(?:\s[^>]*)?>)(.*?)(</w:t>)",
        r"\1\3",
        paragraph_xml,
        flags=re.S,
    )
    xml = re.sub(r"<w:vanish[^>]*/>", "", xml)
    xml = re.sub(r"<w:sz w:val=\"2\"[^>]*/>", "", xml)
    return xml


def _merge_table_xml_preserving_ppr(orig_xml: str, mod_xml: str) -> str:
    """Patch table cell paragraph text only; keep original table/cell/paragraph spacing."""
    o_parts = re.split(r"(<w:p\b.*?</w:p>)", orig_xml, flags=re.S)
    m_paras = re.findall(r"<w:p\b.*?</w:p>", mod_xml, flags=re.S)
    if not m_paras:
        return orig_xml
    para_i = 0
    merged: list[str] = []
    for part in o_parts:
        if part.startswith("<w:p"):
            if para_i < len(m_paras):
                merged.append(_merge_paragraph_xml_preserving_ppr(part, m_paras[para_i]))
                para_i += 1
            else:
                merged.append(part)
        else:
            merged.append(part)
    return "".join(merged)


def _postprocess_paragraph_xml(paragraph_xml: str) -> str:
    """Final layout pass — never remove template spacing from empty paragraphs."""
    text = _paragraph_plain_text_from_xml(paragraph_xml)
    if not text:
        return paragraph_xml
    if _looks_like_role_header_line(text):
        return _strip_italic_from_paragraph_xml(paragraph_xml)
    return paragraph_xml


def _postprocess_document_xml(document_xml: str) -> str:
    """Apply spacing and role-header fixes to the merged document.xml."""
    children = _split_document_body_children(document_xml)
    if not children:
        return document_xml
    merged_parts: list[str] = []
    for child in children:
        if child.startswith("<w:p"):
            child = _postprocess_paragraph_xml(child)
        merged_parts.append(child)
    body_match = re.search(r"(<w:body>).*?(</w:body>)", document_xml, re.S)
    if not body_match:
        return document_xml
    merged_body = body_match.group(1) + "".join(merged_parts) + body_match.group(2)
    return document_xml[: body_match.start()] + merged_body + document_xml[body_match.end() :]


def _merge_document_xml_preserving_layout(
    original_xml: str,
    modified_xml: str,
    *,
    editable_paragraph_indices: set[int],
    editable_table_index: int | None,
) -> str | None:
    """
    Keep original page/body XML byte-for-byte for every frozen region; only swap in
    modified summary/skills paragraphs and the work-experience table.
    Unchanged paragraphs inside editable sections keep original XML (column breaks, spacing).
    """
    orig_children = _split_document_body_children(original_xml)
    mod_children = _split_document_body_children(modified_xml)
    if not orig_children or len(orig_children) != len(mod_children):
        return None

    para_idx = 0
    tbl_idx = 0
    merged_parts: list[str] = []
    for o_child, m_child in zip(orig_children, mod_children):
        if o_child.startswith("<w:sectPr"):
            merged_parts.append(o_child)
            continue
        if o_child.startswith("<w:tbl"):
            use_modified = editable_table_index is not None and tbl_idx == editable_table_index
            merged_parts.append(
                _merge_table_xml_preserving_ppr(o_child, m_child) if use_modified else o_child
            )
            tbl_idx += 1
            continue
        if o_child.startswith("<w:p"):
            if para_idx in editable_paragraph_indices:
                merged_parts.append(_merge_paragraph_xml_preserving_ppr(o_child, m_child))
            else:
                merged_parts.append(o_child)
            para_idx += 1
            continue
        merged_parts.append(o_child)

    body_match = re.search(r"(<w:body>).*?(</w:body>)", original_xml, re.S)
    if not body_match:
        return None
    merged_body = body_match.group(1) + "".join(merged_parts) + body_match.group(2)
    result = original_xml[: body_match.start()] + merged_body + original_xml[body_match.end() :]
    return _postprocess_document_xml(result)


def _merge_document_xml_selective_fallback(
    original_xml: str,
    modified_xml: str,
    *,
    editable_paragraph_indices: set[int],
    editable_table_index: int | None,
) -> str | None:
    """
    When python-docx save flattens tables, swap editable paragraphs by ORIGINAL index
    (not sequential counter) so empty/cleared slots do not shift later sections.
    """
    orig_children = _split_document_body_children(original_xml)
    mod_children = _split_document_body_children(modified_xml)
    if not orig_children or not mod_children:
        return None

    mod_paragraphs = [child for child in mod_children if child.startswith("<w:p")]
    mod_tables = [child for child in mod_children if child.startswith("<w:tbl")]

    para_idx = 0
    tbl_idx = 0
    mod_tbl_idx = 0
    merged_parts: list[str] = []

    for child in orig_children:
        if child.startswith("<w:sectPr"):
            merged_parts.append(child)
            continue
        if child.startswith("<w:tbl"):
            if editable_table_index is not None and tbl_idx == editable_table_index and mod_tbl_idx < len(mod_tables):
                merged_parts.append(
                    _merge_table_xml_preserving_ppr(child, mod_tables[mod_tbl_idx])
                )
                mod_tbl_idx += 1
            else:
                merged_parts.append(child)
            tbl_idx += 1
            continue
        if child.startswith("<w:p"):
            if para_idx in editable_paragraph_indices and para_idx < len(mod_paragraphs):
                merged_parts.append(
                    _merge_paragraph_xml_preserving_ppr(child, mod_paragraphs[para_idx])
                )
            else:
                merged_parts.append(child)
            para_idx += 1
            continue
        merged_parts.append(child)

    body_match = re.search(r"(<w:body>).*?(</w:body>)", original_xml, re.S)
    if not body_match:
        return None
    merged_body = body_match.group(1) + "".join(merged_parts) + body_match.group(2)
    result = original_xml[: body_match.start()] + merged_body + original_xml[body_match.end() :]
    return _postprocess_document_xml(result)


def _restore_frozen_docx_parts(
    original_bytes: bytes,
    modified_bytes: bytes,
    *,
    editable_paragraph_indices: set[int] | None = None,
    editable_table_index: int | None = None,
    memory_doc: Document | None = None,
    editable_table_indices: set[int] | None = None,
) -> bytes:
    """
    python-docx save can strip package metadata and subtly alter header drawings.
    Restore header table + media/package parts from the upload so profile photos,
    round crops, fonts, and contact layout stay exactly as uploaded.
    """
    editable_paragraph_indices = editable_paragraph_indices or set()
    editable_table_indices = set(editable_table_indices or ())
    if editable_table_index is not None:
        editable_table_indices.add(editable_table_index)

    with ZipFile(BytesIO(original_bytes)) as orig_zip:
        orig_names = set(orig_zip.namelist())
        orig_document = orig_zip.read("word/document.xml").decode("utf-8")

        preserve_whole = [
            "[Content_Types].xml",
            "_rels/.rels",
            "docProps/core.xml",
            "docProps/app.xml",
            "word/_rels/document.xml.rels",
            "word/styles.xml",
            "word/stylesWithEffects.xml",
            "word/theme/theme1.xml",
            "word/fontTable.xml",
            "word/webSettings.xml",
            "word/settings.xml",
            "word/numbering.xml",
        ]
        preserved: dict[str, bytes] = {}
        for name in preserve_whole:
            if name in orig_names:
                preserved[name] = orig_zip.read(name)
        for name in orig_names:
            if name.startswith("word/media/"):
                preserved[name] = orig_zip.read(name)

    in_buf = BytesIO(modified_bytes)
    out_buf = BytesIO()
    with ZipFile(in_buf, "r") as mod_zip, ZipFile(out_buf, "w") as out_zip:
        for item in mod_zip.infolist():
            data = mod_zip.read(item.filename)
            if item.filename in preserved:
                data = preserved[item.filename]
            elif item.filename == "word/document.xml":
                if memory_doc is not None and not _document_has_textboxes(orig_document):
                    rebuilt = _rebuild_orig_document_selective(
                        orig_document,
                        memory_doc,
                        editable_paragraph_indices=editable_paragraph_indices,
                        editable_table_indices=editable_table_indices,
                    )
                    if rebuilt is not None:
                        data = rebuilt.encode("utf-8")
                    else:
                        mod_document = data.decode("utf-8")
                        orig_table_count = _count_tables_in_document_xml(orig_document)
                        mod_table_count = _count_tables_in_document_xml(mod_document)
                        if _document_has_textboxes(orig_document):
                            data = mod_document.encode("utf-8")
                        elif orig_table_count > 0 and mod_table_count < orig_table_count:
                            fallback = _merge_document_xml_selective_fallback(
                                orig_document,
                                mod_document,
                                editable_paragraph_indices=editable_paragraph_indices,
                                editable_table_index=editable_table_index,
                            )
                            data = (fallback if fallback is not None else orig_document).encode("utf-8")
                        else:
                            merged_document = _merge_document_xml_preserving_layout(
                                orig_document,
                                mod_document,
                                editable_paragraph_indices=editable_paragraph_indices,
                                editable_table_index=editable_table_index,
                            )
                            if merged_document is not None:
                                data = merged_document.encode("utf-8")
                            else:
                                fallback = _merge_document_xml_selective_fallback(
                                    orig_document,
                                    mod_document,
                                    editable_paragraph_indices=editable_paragraph_indices,
                                    editable_table_index=editable_table_index,
                                )
                                data = (fallback if fallback is not None else orig_document).encode("utf-8")
                else:
                    mod_document = data.decode("utf-8")
                    orig_table_count = _count_tables_in_document_xml(orig_document)
                    mod_table_count = _count_tables_in_document_xml(mod_document)
                    if _document_has_textboxes(orig_document):
                        data = mod_document.encode("utf-8")
                    elif orig_table_count > 0 and mod_table_count < orig_table_count:
                        fallback = _merge_document_xml_selective_fallback(
                            orig_document,
                            mod_document,
                            editable_paragraph_indices=editable_paragraph_indices,
                            editable_table_index=editable_table_index,
                        )
                        data = (fallback if fallback is not None else orig_document).encode("utf-8")
                    else:
                        merged_document = _merge_document_xml_preserving_layout(
                            orig_document,
                            mod_document,
                            editable_paragraph_indices=editable_paragraph_indices,
                            editable_table_index=editable_table_index,
                        )
                        if merged_document is not None:
                            data = merged_document.encode("utf-8")
                        else:
                            fallback = _merge_document_xml_selective_fallback(
                                orig_document,
                                mod_document,
                                editable_paragraph_indices=editable_paragraph_indices,
                                editable_table_index=editable_table_index,
                            )
                            data = (fallback if fallback is not None else orig_document).encode("utf-8")
                try:
                    data = _postprocess_document_xml(data.decode("utf-8")).encode("utf-8")
                except Exception:
                    pass
            out_zip.writestr(item, data)

        for name, data in preserved.items():
            if name not in mod_zip.namelist():
                out_zip.writestr(name, data)

    return out_buf.getvalue()


def apply_tailored_sections_to_docx(
    docx_bytes: bytes,
    *,
    contact: str,
    professional_summary: str,
    professional_experience: str,
    skills: str,
    education: str,
    other: str,
    section_header_indices: dict[str, int],
    section_body_indices: dict[str, list[int]],
    contact_paragraph_indices: list[int],
    source_sections: ParsedResume,
    original_filename: str,
    experience_table_rows: list[ExperienceRowRef] | None = None,
    highlight_keywords: list[str] | None = None,
    skills_highlight_keywords: list[str] | None = None,
    experience_bullets_per_role: list[int] | None = None,
    enable_bold: bool = True,
) -> tuple[bytes, str] | None:
    exp_rows = experience_table_rows or []
    highlights = (
        [k.strip() for k in (highlight_keywords or []) if k and k.strip()]
        if enable_bold
        else []
    )
    skill_highlights = (
        [k.strip() for k in (skills_highlight_keywords or []) if k and k.strip()]
        if enable_bold
        else []
    )
    if not _has_editable_docx_targets(section_body_indices, exp_rows):
        return None

    doc = Document(BytesIO(docx_bytes))
    professional_experience = _ensure_tailored_experience_has_bullets(
        professional_experience,
        source_sections.professional_experience,
    )
    fallbacks = {
        "contact": source_sections.contact,
        "professional_summary": source_sections.professional_summary,
        "professional_experience": source_sections.professional_experience,
        "skills": source_sections.skills,
        "education": source_sections.education,
        "other": source_sections.other,
    }
    tailored_by_section = {
        "contact": contact,
        "professional_summary": professional_summary,
        "professional_experience": professional_experience,
        "skills": skills,
        "education": education,
        "other": other,
    }

    for section_name in reversed(_TAILOR_HEADER_SECTIONS):
        if section_name in _FROZEN_DOCX_SECTIONS:
            continue
        elif section_name == "professional_experience":
            exp_indices = section_body_indices.get(section_name, [])
            text = (professional_experience or "").strip() or (fallbacks["professional_experience"] or "").strip()
            if not text:
                continue
            if exp_rows:
                _update_experience_table_rows(
                    doc,
                    exp_rows,
                    text,
                    highlights,
                    enable_bold=enable_bold,
                    bullets_per_role=experience_bullets_per_role,
                )
            elif exp_indices:
                _update_experience_bullets_only(
                    doc,
                    exp_indices,
                    text,
                    highlights,
                    enable_bold=enable_bold,
                    bullets_per_role=experience_bullets_per_role,
                )
            continue
        indices = section_body_indices.get(section_name, [])
        if not indices:
            continue
        text = (tailored_by_section[section_name] or "").strip() or (fallbacks[section_name] or "").strip()
        if not text:
            continue
        _update_section_inplace(
            doc,
            section_name,
            indices,
            text,
            fallbacks[section_name],
            skill_highlights if section_name == "skills" else highlights if section_name == "professional_summary" else None,
            enable_bold=enable_bold,
        )

    out = BytesIO()
    doc.save(out)

    editable_paragraph_indices: set[int] = set(section_body_indices.get("professional_summary", []))
    skill_indices = section_body_indices.get("skills", [])
    editable_paragraph_indices.update(_skill_content_indices(doc, skill_indices))
    editable_table_indices: set[int] = set()
    editable_table_index: int | None = None
    if exp_rows:
        editable_table_index = exp_rows[0].table_idx
        editable_table_indices.add(editable_table_index)
    else:
        editable_paragraph_indices.update(section_body_indices.get("professional_experience", []))

    merged = _restore_frozen_docx_parts(
        docx_bytes,
        out.getvalue(),
        editable_paragraph_indices=editable_paragraph_indices,
        editable_table_index=editable_table_index,
        memory_doc=doc,
        editable_table_indices=editable_table_indices,
    )
    return merged, output_docx_filename(original_filename)


def output_docx_filename(original_filename: str) -> str:
    name = Path(original_filename).name
    name = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", name).strip()
    if not name or name in (".", ".."):
        return "resume-tailored.docx"
    stem = Path(name).stem or "resume"
    return f"{stem}-tailored.docx"
