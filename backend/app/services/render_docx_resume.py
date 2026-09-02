"""Build an ATS-friendly Word resume from the same structured content as the PDF."""
from __future__ import annotations

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor

from app.services.pdf_resume import parse_contact_identity, sanitize_for_pdf
from app.services.rendercv_resume import _education_entries, _skill_entries

ACCENT = "0B5E75"
TEXT = "111827"


def _set_font(run, *, size: float = 9.5, bold: bool = False, italic: bool = False, color: str = TEXT) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def _set_keep_next(paragraph) -> None:
    paragraph._p.get_or_add_pPr().append(OxmlElement("w:keepNext"))


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), ACCENT)
    properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.append(underline)
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Arial")
    fonts.set(qn("w:hAnsi"), "Arial")
    properties.append(fonts)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    properties.append(size)
    run_element.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    _set_keep_next(paragraph)
    run = paragraph.add_run(text.upper())
    _set_font(run, size=11, bold=True, color=ACCENT)
    borders = paragraph._p.get_or_add_pPr().find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        paragraph._p.get_or_add_pPr().append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), ACCENT)
    borders.append(bottom)


def _add_bullet(doc: Document, text: str) -> None:
    clean = sanitize_for_pdf(text or "").strip().lstrip("-*\u2022 ").strip()
    if not clean:
        return
    paragraph = doc.add_paragraph(style="Resume Bullet")
    paragraph.paragraph_format.space_after = Pt(1.2)
    run = paragraph.add_run(clean)
    _set_font(run)


def _add_labeled_line(doc: Document, label: str, details: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1.5)
    label_run = paragraph.add_run(f"{label}: ")
    _set_font(label_run, bold=True)
    details_run = paragraph.add_run(details)
    _set_font(details_run)


def _add_role_header(doc: Document, role: Any) -> None:
    company = sanitize_for_pdf(getattr(role, "company", "") or "").strip()
    title = sanitize_for_pdf(getattr(role, "title", "") or "").strip()
    location = sanitize_for_pdf(getattr(role, "location", "") or "").strip()
    period = sanitize_for_pdf(getattr(role, "period", "") or "").strip()
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_TAB_ALIGNMENT.RIGHT)
    _set_keep_next(paragraph)
    left = " | ".join(part for part in (company, title) if part) or "Professional Experience"
    right = " | ".join(part for part in (location, period) if part)
    left_run = paragraph.add_run(left)
    _set_font(left_run, size=9.7, bold=True)
    if right:
        paragraph.add_run("\t")
        right_run = paragraph.add_run(right)
        _set_font(right_run, size=9.2)


def _add_education(doc: Document, education: str) -> None:
    _add_education_entries(doc, _education_entries(education))


def _add_education_entries(doc: Document, entries: list[dict[str, Any]]) -> None:
    for entry in entries:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.15), WD_TAB_ALIGNMENT.RIGHT)
        _set_keep_next(paragraph)
        institution = str(entry.get("institution") or "Education").strip()
        degree = str(entry.get("degree") or "").strip()
        area = str(entry.get("area") or "").strip()
        left_run = paragraph.add_run(institution)
        _set_font(left_run, bold=True)
        detail = " in ".join(part for part in (degree, area) if part) if degree else area
        if detail:
            detail_run = paragraph.add_run(f", {detail}")
            _set_font(detail_run)
        right = " | ".join(
            str(entry.get(key) or "").strip() for key in ("location", "date") if entry.get(key)
        )
        if right:
            paragraph.add_run("\t")
            right_run = paragraph.add_run(right)
            _set_font(right_run, size=9.2)
        for highlight in entry.get("highlights", []) or []:
            _add_bullet(doc, str(highlight))


def build_docx_resume(
    *,
    contact: str,
    professional_summary: str,
    roles: list[Any],
    bullets_by_role: list[list[str]],
    skills: str,
    education: str,
    other: str,
    skill_entries: list[dict[str, str]] | None = None,
    education_entries: list[dict[str, Any]] | None = None,
    extra_sections: list[tuple[str, list[str]]] | None = None,
) -> bytes:
    """Return a polished single-column DOCX containing the final generated resume.

    ``skill_entries`` / ``education_entries`` / ``extra_sections``, when supplied,
    bypass the string re-parsing and render the typed data straight through.
    """
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    bullet_style = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.base_style = doc.styles["List Bullet"]
    bullet_style.font.name = "Arial"
    bullet_style.font.size = Pt(9.5)
    bullet_style.paragraph_format.left_indent = Inches(0.22)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.16)

    identity = parse_contact_identity(contact or "")
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(1)
    name_run = name.add_run(identity.name or "Candidate")
    _set_font(name_run, size=20, bold=True, color=ACCENT)
    if identity.headline:
        headline = doc.add_paragraph()
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        headline.paragraph_format.space_after = Pt(2.5)
        headline_run = headline.add_run(identity.headline)
        _set_font(headline_run, size=10.5, bold=True, color=ACCENT)

    contact_line = doc.add_paragraph()
    contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_line.paragraph_format.space_after = Pt(4)
    items: list[tuple[str, str | None]] = []
    if identity.email:
        items.append((identity.email, f"mailto:{identity.email}"))
    if identity.phone:
        phone_href = re.sub(r"[^+\d]", "", identity.phone)
        items.append((identity.phone, f"tel:{phone_href}"))
    if identity.location:
        items.append((identity.location, None))
    if identity.linkedin_url:
        items.append(("LinkedIn", identity.linkedin_url))
    if identity.portfolio_url:
        items.append(("Portfolio", identity.portfolio_url))
    if identity.github_url:
        items.append(("GitHub", identity.github_url))
    items.extend((label or "Website", url) for label, url in identity.other_links)
    for index, (label, url) in enumerate(items):
        if index:
            separator = contact_line.add_run("  |  ")
            _set_font(separator, size=9)
        if url:
            _add_hyperlink(contact_line, label, url)
        else:
            detail_run = contact_line.add_run(label)
            _set_font(detail_run, size=9)

    if professional_summary.strip():
        _add_section_heading(doc, "Professional Summary")
        for block in re.split(r"\n\s*\n|\n", sanitize_for_pdf(professional_summary).strip()):
            if not block.strip():
                continue
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(block.strip())
            _set_font(run)

    resolved_skill_entries = skill_entries if skill_entries is not None else _skill_entries(skills)
    if resolved_skill_entries:
        _add_section_heading(doc, "Skills")
        for item in resolved_skill_entries:
            _add_labeled_line(doc, item["label"], item["details"])

    if roles:
        _add_section_heading(doc, "Professional Experience")
        for index, role in enumerate(roles):
            _add_role_header(doc, role)
            for bullet in bullets_by_role[index] if index < len(bullets_by_role) else []:
                _add_bullet(doc, bullet)

    if education_entries is not None:
        if education_entries:
            _add_section_heading(doc, "Education")
            _add_education_entries(doc, education_entries)
    elif education.strip():
        _add_section_heading(doc, "Education")
        _add_education(doc, education)

    if extra_sections is not None:
        for heading, entries in extra_sections:
            clean = [e.strip() for e in entries if e and e.strip()]
            if not clean:
                continue
            _add_section_heading(doc, heading or "Additional")
            for line in clean:
                _add_bullet(doc, line)
    else:
        other_lines = [line.strip() for line in sanitize_for_pdf(other or "").splitlines() if line.strip()]
        if other_lines:
            _add_section_heading(doc, "Additional")
            for line in other_lines:
                if re.match(r"^[\-*\u2022]\s+", line):
                    _add_bullet(doc, line)
                else:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    _set_font(run)

    properties = doc.core_properties
    properties.title = f"{identity.name or 'Candidate'} Resume"
    properties.subject = "Tailored Resume"
    properties.author = identity.name or "Candidate"
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def build_docx_resume_from_model(model: Any) -> bytes:
    """Typed twin of :func:`build_docx_resume` — renders a ``ResumeModel`` directly."""
    from app.services.rendercv_resume import (
        _education_entries_from_model,
        _skill_entries_from_model,
    )

    return build_docx_resume(
        contact=model.contact_block(),
        professional_summary=model.professional_summary,
        roles=model.work_experience_roles(),
        bullets_by_role=[list(role.responsibilities) for role in model.professional_experience],
        skills="",
        education="",
        other="",
        skill_entries=_skill_entries_from_model(model.technical_skills),
        education_entries=_education_entries_from_model(model.education),
        extra_sections=[(heading, list(entries)) for heading, entries in model.extra_sections()],
    )
