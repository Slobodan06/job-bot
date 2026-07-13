"""Compare original Ciro docx spacing vs tailored PDF (styling only)."""
from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path

from docx import Document

from app.services.docx_resume import (
    _all_document_paragraphs,
    _is_bullet_paragraph,
    _paragraph_text,
    parse_resume_from_docx,
)
from app.services.extract_text import extract_text_from_bytes

DOCX = Path(r"C:\Users\Jelix\Downloads\cv_ciro-fullstack (3).docx")
PDF = Path(r"C:\Users\Jelix\Downloads\cv_ciro-fullstack (3)-tailored.pdf")


def spacing_from_xml(xml: str) -> tuple[str, str]:
    before_m = re.search(r'w:before="(\d+)"', xml)
    after_m = re.search(r'w:after="(\d+)"', xml)
    return (before_m.group(1) if before_m else "0", after_m.group(1) if after_m else "0")


def main() -> None:
    docx_bytes = DOCX.read_bytes()
    parsed_doc = parse_resume_from_docx(docx_bytes)
    d = Document(BytesIO(docx_bytes))
    paras = _all_document_paragraphs(d)

    print("=== ORIGINAL DOCX STRUCTURE ===")
    print("detected_role_count:", parsed_doc.detected_role_count)
    print("experience_table_rows:", len(parsed_doc.experience_table_rows))
    print("skills paragraph indices:", len(parsed_doc.section_body_indices.get("skills", [])))

    skill_idx = parsed_doc.section_body_indices.get("skills", [])
    print("\n--- SKILLS (paragraph index, empty, bullet, spacing before/after, text) ---")
    for i in skill_idx:
        if i >= len(paras):
            continue
        p = paras[i]
        t = _paragraph_text(p)
        b, a = spacing_from_xml(p._element.xml)
        print(
            f"  [{i:3d}] empty={not t.strip():5} bullet={str(_is_bullet_paragraph(p)):5} "
            f"spc={b}/{a} | {t[:70]!r}"
        )

    if parsed_doc.experience_table_rows:
        print("\n--- EXPERIENCE TABLE (per row, paragraphs in content cell) ---")
        for ri, ref in enumerate(parsed_doc.experience_table_rows):
            cell = d.tables[ref.table_idx].rows[ref.row_idx].cells[ref.content_cols[0]]
            print(f"\n  Row {ri + 1} ({len(cell.paragraphs)} paragraphs):")
            for pi, para in enumerate(cell.paragraphs):
                t = _paragraph_text(para).strip()
                b, a = spacing_from_xml(para._element.xml)
                kind = "BULLET" if _is_bullet_paragraph(para) else ("EMPTY" if not t else "HEADER")
                print(f"    p{pi:2d} {kind:6s} spc={b}/{a} | {t[:72]!r}")

    pdf_text = extract_text_from_bytes(PDF.name, PDF.read_bytes())
    lines = pdf_text.splitlines()
    print("\n=== TAILORED PDF (experience area, non-empty lines) ===")
    in_exp = False
    shown = 0
    blank_run = 0
    for line in lines:
        stripped = line.strip()
        if re.search(r"^(WORK\s+)?EXPERIENCE", stripped, re.I):
            in_exp = True
            print(f"  [SECTION] {stripped}")
            continue
        if in_exp and re.search(r"^EDUCATION", stripped, re.I):
            print(f"  [SECTION END] {stripped}")
            break
        if not in_exp:
            continue
        if not stripped:
            blank_run += 1
            if blank_run == 1:
                print("  --- blank line ---")
            continue
        blank_run = 0
        print(f"  {stripped[:88]}")
        shown += 1
        if shown > 45:
            print("  ...")
            break

    # Count blank lines between job titles in PDF
    print("\n=== PDF: blank lines between consecutive job-title-like lines ===")
    title_re = re.compile(
        r"\b(developer|engineer|manager|architect|consultant|specialist|director|lead)\b",
        re.I,
    )
    prev_title_idx = None
    for i, line in enumerate(lines):
        s = line.strip()
        if not s or not title_re.search(s) or len(s) > 80:
            continue
        if prev_title_idx is not None:
            between = lines[prev_title_idx + 1 : i]
            blank_count = sum(1 for b in between if not b.strip())
            content_between = [b.strip() for b in between if b.strip()]
            print(
                f"  Gap before {s[:45]!r}: {blank_count} blank lines, "
                f"{len(content_between)} content lines between titles"
            )
        prev_title_idx = i


if __name__ == "__main__":
    main()
