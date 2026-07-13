from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"D:\work\Job-bot\job-bot\Daniel_Wiseman_AI_Engineer_HRA.docx"

NAVY = RGBColor(20, 45, 70)
GRAY = RGBColor(75, 75, 75)

def set_font(run, size=10.2, bold=False, color=None, italic=False):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def set_repeat_keep(p, keep_next=False, keep_lines=True):
    pPr = p._p.get_or_add_pPr()
    if keep_next:
        pPr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        pPr.append(OxmlElement("w:keepLines"))

def add_section(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    r = p.add_run(title.upper())
    set_font(r, 11.4, True, NAVY)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "143A5A")
    borders.append(bottom)
    pPr.append(borders)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    p.paragraph_format.space_after = Pt(2.0)
    p.paragraph_format.line_spacing = 1.0
    set_repeat_keep(p)
    set_font(p.add_run(text), 9.7)
    return p

def add_role(doc, company, title, dates, location=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.35))
    set_font(p.add_run(company), 10.2, True, NAVY)
    set_font(p.add_run(f" | {title}"), 10.2, True)
    if location:
        set_font(p.add_run(f" | {location}"), 9.5, False, GRAY)
    set_font(p.add_run("\t" + dates), 9.6, True, GRAY)
    p.paragraph_format.keep_with_next = True

doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.48)
sec.bottom_margin = Inches(0.48)
sec.left_margin = Inches(0.55)
sec.right_margin = Inches(0.55)
sec.header_distance = Inches(0.25)
sec.footer_distance = Inches(0.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.2)
normal.paragraph_format.space_after = Pt(2)
normal.paragraph_format.line_spacing = 1.02

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(1)
set_font(title.add_run("DANIEL WISEMAN"), 19, True, NAVY)

tag = doc.add_paragraph()
tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
tag.paragraph_format.space_after = Pt(2)
set_font(tag.add_run("AI ENGINEER | PYTHON, LLM APIs, DOCUMENT AUTOMATION & HEALTHCARE SYSTEMS"), 10.1, True)

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact.paragraph_format.space_after = Pt(4)
set_font(contact.add_run("San Francisco, CA | (818) 863-6443 | daniel.wiseman.ca@gmail.com"), 9.3, color=GRAY)

add_section(doc, "Professional Summary")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(3)
set_font(p.add_run(
    "AI-focused software engineer with 10+ years building production systems across healthcare, enterprise SaaS, and infrastructure. Hands-on experience developing Python services, OpenAI API integrations, relational data workflows, and AI document automation that converts unstructured PDF/DOCX inputs into structured, validated, user-ready documents. Built evidence-grounded LLM guardrails, structured JSON pipelines, and deterministic output generation using API patterns directly transferable to the Claude API. Brings HIPAA-aware engineering judgment and a record of improving performance, reliability, and release speed in systems used by millions."
), 9.8)

add_section(doc, "Core Skills")
skills = [
    ("LLM & AI Engineering", "OpenAI API, LLM API integration, prompt design, structured JSON output, multi-block response handling, output parsing, hallucination controls, source-grounded validation; Claude API patterns transferable from comparable LLM APIs"),
    ("Document Intelligence", "Unstructured-to-structured document conversion, PDF/DOCX ingestion, text extraction, document section detection, data normalization, DOCX/PDF generation, parse-failure handling, NLP, data preprocessing"),
    ("Programming & APIs", "Python, FastAPI, Django, Flask, Node.js, REST APIs, GraphQL, asynchronous services, microservices"),
    ("Data", "SQL, Microsoft SQL Server, PostgreSQL, MySQL, MongoDB, database design, multi-table queries, query optimization"),
    ("Reliability & Delivery", "Pytest, automated testing, guardrails, audit logging, monitoring, Git, GitHub Actions, CI/CD, Docker, Kubernetes, Terraform"),
    ("Cloud & Regulated Systems", "Azure, AWS, HIPAA, PHI/PII security, healthcare systems, FHIR, HL7, EHR and clinical workflows"),
]
for label, value in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1.2)
    set_font(p.add_run(label + ": "), 9.4, True, NAVY)
    set_font(p.add_run(value), 9.4)

add_section(doc, "Selected AI Engineering Project")
add_role(doc, "Job Bot", "AI Engineer / Full-Stack Developer", "2026", "Independent Project")
add_bullet(doc, "Built a production-style FastAPI and React application that accepts unstructured PDF and DOCX resumes, combines them with job descriptions, and uses the OpenAI API to generate ATS-targeted content.")
add_bullet(doc, "Designed prompts and structured JSON contracts for resume and cover-letter workflows, converting free-form source documents into predictable summaries, skills, experience bullets, qualification gaps, and application documents.")
add_bullet(doc, "Implemented asynchronous LLM API calls, JSON response parsing, malformed-response handling, sanitized error reporting, and deterministic fallbacks when the model or API is unavailable.")
add_bullet(doc, "Engineered evidence-first guardrails that map generated claims to source facts, reject unsupported technologies and metrics, and prevent an LLM from turning adjacent experience into confidently false statements.")
add_bullet(doc, "Validated generated output against the uploaded source before export, preserving employer boundaries, employment chronology, verified metrics, contact details, and education while flagging unsupported job requirements.")
add_bullet(doc, "Built document-processing pipelines for PDF/DOCX text extraction, section detection, content normalization, in-place Word updates, fresh document generation, and DOCX-to-PDF conversion.")
add_bullet(doc, "Handled incorrect or incomplete parses through file-type validation, extraction errors, section-aware fallbacks, frozen source sections, and source-evidence checks before a document reaches the final export path.")
add_bullet(doc, "Created 40 automated tests covering evidence validation, unsupported-claim filtering, skill transfer logic, metric preservation, structured document behavior, and adversarial LLM-output scenarios.")
add_bullet(doc, "Containerized the full-stack service with Docker and configured production deployment, LibreOffice-based document conversion, health checks, and explicit reporting of whether the LLM was used.")

add_section(doc, "Professional Experience")
add_role(doc, "CVS Health", "Senior Software Engineer", "01/2022 - 03/2026", "Remote")
add_bullet(doc, "Built and maintained production services for a unified digital health platform integrating pharmacy, insurance, and care workflows for 15M+ registered members.")
add_bullet(doc, "Engineered Node.js microservices and API gateways for prescription fulfillment, eligibility, and identity workflows, improving peak-load response times by 29%.")
add_bullet(doc, "Improved sustained concurrent-user capacity by 34% through targeted caching and backend concurrency work; maintained 99.9% message-delivery reliability across Kafka pipelines.")
add_bullet(doc, "Reduced deployment lead time by 78% using Docker, Kubernetes, Terraform, and automated test/release pipelines while improving release stability.")
add_bullet(doc, "Established HIPAA-aligned access controls, encrypted service communication, and audit logging for systems handling patient and member data.")

add_role(doc, "Dignity Health", "Senior Software Engineer", "08/2018 - 12/2021")
add_bullet(doc, "Developed patient-facing and internal healthcare platforms supporting clinical, pharmacy, and retail workflows across hospitals serving millions of patients annually.")
add_bullet(doc, "Designed Node.js APIs and event-driven integrations across EHR, prescription, and retail systems, increasing peak transaction throughput by 36% and reducing cross-system data inconsistencies by 33%.")
add_bullet(doc, "Tuned multi-system queries and service contracts across PostgreSQL and Microsoft SQL Server, improving response times by 35% for critical clinical and pharmacy endpoints.")
add_bullet(doc, "Standardized automated testing and release pipelines, lowering production incidents by 42% while increasing deployment frequency and reliability.")
add_bullet(doc, "Applied HIPAA-aligned access control, auditing, and data-protection practices to applications handling PHI and clinical information.")

add_role(doc, "Meter", "Backend Engineer", "06/2015 - 07/2018")
add_bullet(doc, "Built production backend services and REST APIs for a cloud-managed networking platform supporting secure provisioning and management of 2,000+ devices.")
add_bullet(doc, "Developed Kafka telemetry pipelines ingesting millions of network events daily, improving observability and troubleshooting efficiency by 41%.")
add_bullet(doc, "Created authentication, authorization, and configuration-validation services that reduced misconfiguration-related outages by 32%.")
add_bullet(doc, "Operated containerized AWS services with Terraform, sustaining 99.9% uptime during continuous delivery and customer expansion.")

add_section(doc, "Education")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
set_font(p.add_run("California Polytechnic State University, San Luis Obispo"), 9.8, True, NAVY)
set_font(p.add_run(" | Bachelor of Science, Computer Science | 05/2015 | GPA: 3.63"), 9.6)

# Quiet page-number footer for a multi-page professional resume.
footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Daniel Wiseman | AI Engineer")
set_font(run, 8.2, color=GRAY)

doc.core_properties.title = "Daniel Wiseman - AI Engineer Resume"
doc.core_properties.subject = "Application for AI Engineer at Healthcare Retroactive Audits"
doc.core_properties.author = "Daniel Wiseman"
doc.core_properties.keywords = "AI Engineer, Python, SQL Server, OpenAI API, LLM, PDF, DOCX, HIPAA"
doc.save(OUT)
print(OUT)
